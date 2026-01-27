from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File, Request
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import uuid
from pathlib import Path
from typing import List, Optional
from datetime import datetime, timedelta
import io
import zipfile
import json
import pandas as pd

# Import models
from models import (
    User, UserCreate, UserLogin, Token,
    Question, QuestionCreate, QuestionCategory, DifficultyLevel,
    UserProgress, UserAnswer, CategoryProgress, StudySession,
    AIGenerationRequest, AIGenerationResponse,
    ExamSession, ExamResult, QuestionReport,
    UserAnalytics, DataExportRequest,
    APIToken, APITokenCreate, APITokenResponse, APITokenScope, SetupTokenRequest
)

# Import services
from auth_service import (
    hash_password, verify_password, create_access_token,
    create_refresh_token, get_current_user, create_api_token_model,
    verify_api_token
)
from storage_service import storage_service
from ai_service import ai_service
from adaptive_learning import adaptive_engine

# Import email service
from email_service import (
    send_verification_email,
    send_password_reset_email,
    send_qualifying_session_email,
    send_email_from_support,
    send_question_report_notification,
    send_contact_form_notification
)

# Import email client service (for admin email client)
from email_client_service import (
    fetch_emails,
    fetch_email_by_id,
    send_email_from_support as send_support_email,
    get_unread_count
)

# Import tenant configuration
from tenant_config import (
    TenantConfig, TenantCreate, TenantUpdate, DEFAULT_TENANT,
    get_tenant_by_domain, get_tenant_by_id, get_all_tenants,
    create_tenant, update_tenant, deactivate_tenant, ensure_default_tenant,
    clear_tenant_cache
)
from tenant_middleware import get_current_tenant, get_tenant_config

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app
app = FastAPI(title="MedMCQ API", version="1.0.0")

# Store db in app state for middleware access
app.state.db = db

# Create API router
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Root-level health check for deployment (MUST be before api_router)
@app.get("/health")
async def root_health_check():
    """Root-level health check endpoint for Kubernetes deployment."""
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}


# Startup event for tenant initialization and index creation
@app.on_event("startup")
async def startup_event():
    """Initialize tenant system and create database indexes on startup."""
    logger.info("Starting MedMCQ API - Initializing multi-tenant system...")
    
    # Ensure default tenant exists
    await ensure_default_tenant(db)
    
    # Create indexes for tenant queries
    try:
        # Tenants collection indexes
        await db.tenants.create_index("domain", unique=True)
        await db.tenants.create_index("tenant_id", unique=True)
        logger.info("Created indexes on tenants collection")
    except Exception as e:
        logger.warning(f"Tenant indexes may already exist: {e}")
    
    try:
        # Users - lookup by email within tenant
        await db.users.create_index([("tenant_id", 1), ("email", 1)])
        logger.info("Created compound index on users (tenant_id, email)")
    except Exception as e:
        logger.warning(f"User index may already exist: {e}")
    
    try:
        # Questions - filter by tenant and category
        await db.questions.create_index([("tenant_id", 1), ("category", 1)])
        await db.questions.create_index([("tenant_id", 1), ("source", 1)])
        logger.info("Created indexes on questions collection")
    except Exception as e:
        logger.warning(f"Question indexes may already exist: {e}")
    
    try:
        # User progress - lookup by user and tenant
        await db.user_progress.create_index([("tenant_id", 1), ("user_id", 1)])
        logger.info("Created index on user_progress collection")
    except Exception as e:
        logger.warning(f"User progress index may already exist: {e}")
    
    try:
        # Study sessions - lookup by user, tenant, date
        await db.study_sessions.create_index([("tenant_id", 1), ("user_id", 1), ("date", 1)])
        logger.info("Created index on study_sessions collection")
    except Exception as e:
        logger.warning(f"Study sessions index may already exist: {e}")
    
    try:
        # Exam sessions - lookup by tenant and user
        await db.exam_sessions.create_index([("tenant_id", 1), ("user_id", 1)])
        logger.info("Created index on exam_sessions collection")
    except Exception as e:
        logger.warning(f"Exam sessions index may already exist: {e}")
    
    logger.info("Multi-tenant system initialization complete")

# ============================================================================
# TENANT CONFIGURATION ROUTES
# ============================================================================

@api_router.get("/tenant/config", response_model=TenantConfig)
async def get_tenant_configuration(
    request: Request,
    domain: Optional[str] = None
):
    """
    Get tenant configuration.
    
    This is a PUBLIC endpoint (no authentication required) so the frontend
    can fetch tenant branding before user logs in.
    
    Args:
        domain: Optional domain to look up. If not provided, uses request Host header.
    
    Returns:
        TenantConfig for the specified or detected domain
    """
    if domain:
        # Look up by provided domain
        tenant = await get_tenant_by_domain(db, domain)
        if tenant:
            return tenant
    else:
        # Use the middleware to detect from request
        tenant_config = await get_tenant_config(request)
        return tenant_config
    
    # Fallback to default
    return DEFAULT_TENANT


# ============================================================================
# ADMIN TENANT MANAGEMENT ROUTES
# ============================================================================

@api_router.post("/admin/tenants", response_model=TenantConfig)
async def admin_create_tenant(
    tenant_data: TenantCreate,
    user_id: str = Depends(get_current_user)
):
    """Create a new tenant (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Check if tenant_id or domain already exists
    existing_by_id = await get_tenant_by_id(db, tenant_data.tenant_id)
    if existing_by_id:
        raise HTTPException(
            status_code=400,
            detail=f"Tenant with ID '{tenant_data.tenant_id}' already exists"
        )
    
    existing_by_domain = await get_tenant_by_domain(db, tenant_data.domain)
    if existing_by_domain:
        raise HTTPException(
            status_code=400,
            detail=f"Tenant with domain '{tenant_data.domain}' already exists"
        )
    
    tenant = await create_tenant(db, tenant_data)
    logger.info(f"Admin {user_id} created tenant: {tenant.tenant_id}")
    return tenant


@api_router.put("/admin/tenants/{tenant_id}", response_model=TenantConfig)
async def admin_update_tenant(
    tenant_id: str,
    update_data: TenantUpdate,
    user_id: str = Depends(get_current_user)
):
    """Update a tenant configuration (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Check if domain is being changed to an existing one
    if update_data.domain:
        existing = await get_tenant_by_domain(db, update_data.domain)
        if existing and existing.tenant_id != tenant_id:
            raise HTTPException(
                status_code=400,
                detail=f"Domain '{update_data.domain}' is already in use"
            )
    
    updated = await update_tenant(db, tenant_id, update_data)
    if not updated:
        raise HTTPException(status_code=404, detail="Tenant not found")
    
    logger.info(f"Admin {user_id} updated tenant: {tenant_id}")
    return updated


@api_router.get("/admin/tenants", response_model=List[TenantConfig])
async def admin_list_tenants(
    include_inactive: bool = False,
    user_id: str = Depends(get_current_user)
):
    """List all tenants (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    return await get_all_tenants(db, include_inactive)


@api_router.delete("/admin/tenants/{tenant_id}")
async def admin_deactivate_tenant(
    tenant_id: str,
    user_id: str = Depends(get_current_user)
):
    """Deactivate a tenant (admin only). This is a soft delete."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Prevent deactivating the default tenant
    if tenant_id == "med":
        raise HTTPException(
            status_code=400,
            detail="Cannot deactivate the default tenant"
        )
    
    success = await deactivate_tenant(db, tenant_id)
    if not success:
        raise HTTPException(status_code=404, detail="Tenant not found")
    
    logger.info(f"Admin {user_id} deactivated tenant: {tenant_id}")
    return {"success": True, "message": f"Tenant '{tenant_id}' has been deactivated"}


@api_router.post("/admin/run-migration")
async def admin_run_migration(user_id: str = Depends(get_current_user)):
    """Run data migration to add tenant_id to existing records (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Import and run migration
    from migrations.add_tenant_id import run_migration
    result = await run_migration()
    
    logger.info(f"Admin {user_id} ran migration: {result}")
    return result


# ============================================================================
# ADMIN REPORTED ISSUES ROUTES
# ============================================================================

@api_router.get("/admin/reported-issues")
async def admin_get_reported_issues(user_id: str = Depends(get_current_user)):
    """Get all reported question issues (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Get all reports with question details
    reports = await db.question_reports.find({}, {"_id": 0}).sort("timestamp", -1).to_list(500)
    
    # Enrich with question data
    for report in reports:
        question = await db.questions.find_one(
            {"id": report.get("question_id")}, 
            {"_id": 0, "question": 1, "category": 1, "quarantined": 1}
        )
        if question:
            report["question_text"] = question.get("question", "N/A")
            report["question_category"] = question.get("category", "N/A")
            report["question_quarantined"] = question.get("quarantined", False)
        
        # Get reporter info
        user = await db.users.find_one(
            {"id": report.get("user_id")},
            {"_id": 0, "email": 1, "full_name": 1}
        )
        if user:
            report["reporter_email"] = user.get("email", "Unknown")
            report["reporter_name"] = user.get("full_name", "Unknown")
    
    return reports


@api_router.put("/admin/reported-issues/{question_id}/status")
async def admin_update_report_status(
    question_id: str,
    data: dict,
    user_id: str = Depends(get_current_user)
):
    """Update the status of a reported question (admin only).
    
    Status options: pending, fixed, quarantined
    """
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    new_status = data.get("status")
    if new_status not in ["pending", "fixed", "quarantined"]:
        raise HTTPException(status_code=400, detail="Invalid status. Use: pending, fixed, quarantined")
    
    # Update all reports for this question
    await db.question_reports.update_many(
        {"question_id": question_id},
        {"$set": {"status": new_status, "resolved": new_status in ["fixed", "quarantined"]}}
    )
    
    # Update question quarantine status
    if new_status == "quarantined":
        await db.questions.update_one(
            {"id": question_id},
            {"$set": {"quarantined": True}}
        )
    elif new_status == "fixed":
        await db.questions.update_one(
            {"id": question_id},
            {"$set": {"quarantined": False}}
        )
    
    logger.info(f"Admin {user_id} updated question {question_id} status to {new_status}")
    return {"success": True, "message": f"Question status updated to {new_status}"}


# ============================================================================
# ADMIN EMAIL ROUTES
# ============================================================================

@api_router.get("/admin/users-for-email")
async def admin_get_users_for_email(user_id: str = Depends(get_current_user)):
    """Get list of users for email sending (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    users = await db.users.find(
        {},
        {"_id": 0, "id": 1, "email": 1, "full_name": 1, "subscription_plan": 1, "created_at": 1}
    ).sort("created_at", -1).to_list(1000)
    
    return users


# ============================================================================
# ADMIN EMAIL CLIENT ROUTES (Full IMAP/SMTP client for support@)
# ============================================================================

@api_router.get("/admin/email/inbox")
async def admin_get_inbox(
    limit: int = 50,
    user_id: str = Depends(get_current_user)
):
    """Get inbox emails from support@medmcq.com.au (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    emails = fetch_emails(folder="INBOX", limit=limit)
    return {"emails": emails, "folder": "inbox", "count": len(emails)}


@api_router.get("/admin/email/sent")
async def admin_get_sent(
    limit: int = 50,
    user_id: str = Depends(get_current_user)
):
    """Get sent emails from support@medmcq.com.au (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    emails = fetch_emails(folder="Sent", limit=limit)
    return {"emails": emails, "folder": "sent", "count": len(emails)}


@api_router.get("/admin/email/unread-count")
async def admin_get_unread_count(user_id: str = Depends(get_current_user)):
    """Get unread email count (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    count = get_unread_count()
    return {"unread_count": count}


@api_router.get("/admin/email/message/{email_id}")
async def admin_get_email_message(
    email_id: str,
    folder: str = "INBOX",
    user_id: str = Depends(get_current_user)
):
    """Get full email message by ID (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    email_data = fetch_email_by_id(email_id, folder)
    
    if not email_data:
        raise HTTPException(status_code=404, detail="Email not found")
    
    return email_data


@api_router.post("/admin/email/send")
async def admin_send_email(
    data: dict,
    user_id: str = Depends(get_current_user)
):
    """Send email from support@medmcq.com.au (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    to_email = data.get("to")
    subject = data.get("subject")
    body = data.get("body")
    reply_to_message_id = data.get("reply_to_message_id")
    
    if not to_email or not subject or not body:
        raise HTTPException(status_code=400, detail="to, subject, and body are required")
    
    success = send_support_email(
        to_email=to_email,
        subject=subject,
        body=body,
        reply_to_message_id=reply_to_message_id
    )
    
    if success:
        # Log the sent email
        await db.admin_email_log.insert_one({
            "admin_id": user_id,
            "to_email": to_email,
            "subject": subject,
            "body": body,
            "sent_at": datetime.utcnow().isoformat(),
            "success": True
        })
        return {"success": True, "message": "Email sent successfully"}
    else:
        raise HTTPException(status_code=500, detail="Failed to send email")


@api_router.post("/admin/email/reply")
async def admin_reply_to_email(
    data: dict,
    user_id: str = Depends(get_current_user)
):
    """Reply to an email (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    original_email_id = data.get("original_email_id")
    folder = data.get("folder", "INBOX")
    body = data.get("body")
    
    if not original_email_id or not body:
        raise HTTPException(status_code=400, detail="original_email_id and body are required")
    
    # Fetch original email to get details
    original = fetch_email_by_id(original_email_id, folder)
    
    if not original:
        raise HTTPException(status_code=404, detail="Original email not found")
    
    # Construct reply
    to_email = original["from"]["email"]
    subject = original["subject"]
    if not subject.lower().startswith("re:"):
        subject = f"Re: {subject}"
    
    # Add quoted original message
    reply_body = f"{body}\n\n---\nOn {original['date']}, {original['from']['name']} <{original['from']['email']}> wrote:\n\n{original['body']}"
    
    success = send_support_email(
        to_email=to_email,
        subject=subject,
        body=reply_body,
        reply_to_message_id=original.get("message_id")
    )
    
    if success:
        await db.admin_email_log.insert_one({
            "admin_id": user_id,
            "to_email": to_email,
            "subject": subject,
            "body": reply_body,
            "reply_to": original_email_id,
            "sent_at": datetime.utcnow().isoformat(),
            "success": True
        })
        return {"success": True, "message": "Reply sent successfully"}
    else:
        raise HTTPException(status_code=500, detail="Failed to send reply")


@api_router.get("/admin/dashboard/multi-tenant")
async def admin_multi_tenant_dashboard(user_id: str = Depends(get_current_user)):
    """Get multi-tenant dashboard statistics (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Get all tenants
    tenants = await get_all_tenants(db, include_inactive=True)
    
    # Aggregate stats per tenant
    tenant_stats = []
    for tenant in tenants:
        tid = tenant.tenant_id
        
        # Count users for this tenant
        user_count = await db.users.count_documents({"tenant_id": tid})
        
        # Count questions for this tenant
        question_count = await db.questions.count_documents({"tenant_id": tid})
        
        # Count active subscriptions (non-free)
        active_subs = await db.users.count_documents({
            "tenant_id": tid,
            "subscription_status": {"$in": ["active", "free_grant"]},
            "subscription_plan": {"$ne": "free"}
        })
        
        # Count study sessions in last 30 days
        thirty_days_ago = (datetime.utcnow() - timedelta(days=30)).date().isoformat()
        recent_sessions = await db.study_sessions.count_documents({
            "tenant_id": tid,
            "date": {"$gte": thirty_days_ago}
        })
        
        tenant_stats.append({
            "tenant_id": tid,
            "name": tenant.name,
            "domain": tenant.domain,
            "is_active": tenant.is_active,
            "users": user_count,
            "questions": question_count,
            "active_subscriptions": active_subs,
            "recent_sessions_30d": recent_sessions
        })
    
    # Overall totals
    total_users = sum(t["users"] for t in tenant_stats)
    total_questions = sum(t["questions"] for t in tenant_stats)
    total_subscriptions = sum(t["active_subscriptions"] for t in tenant_stats)
    
    return {
        "tenants": tenant_stats,
        "totals": {
            "total_tenants": len(tenants),
            "active_tenants": len([t for t in tenants if t.is_active]),
            "total_users": total_users,
            "total_questions": total_questions,
            "total_active_subscriptions": total_subscriptions
        }
    }

# ============================================================================
# AUTHENTICATION ROUTES (Tenant-Aware)
# ============================================================================

@api_router.post("/auth/register", response_model=User)
async def register(
    user_data: UserCreate,
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Register a new user with tenant awareness.
    
    Users are registered to the tenant determined by the request domain.
    """
    # Check if user exists within this tenant
    existing_user = await db.users.find_one({
        "email": user_data.email,
        "tenant_id": tenant_id
    })
    if existing_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Email already registered"
        )
    
    # Free users MUST accept marketing consent
    if not user_data.marketing_consent:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="You must accept the marketing communications disclaimer to create a free account"
        )
    
    # Generate email verification token
    verification_token = str(uuid.uuid4())
    
    # Create user with tenant_id
    user = User(
        email=user_data.email,
        full_name=user_data.full_name,
        institution=user_data.institution,
        current_year=user_data.current_year,
        degree_type=user_data.degree_type,
        country=user_data.country,
        marketing_consent=user_data.marketing_consent,
        email_verified=False,
        verification_token=verification_token,
        tenant_id=tenant_id  # Multi-tenant: assign to current tenant
    )
    
    # Hash password and store separately
    hashed_password = hash_password(user_data.password)
    
    # Store user in database
    user_dict = user.model_dump()
    user_dict['hashed_password'] = hashed_password
    user_dict['created_at'] = user_dict['created_at'].isoformat()
    
    await db.users.insert_one(user_dict)
    
    # Initialize user progress with tenant_id
    progress = UserProgress(user_id=user.id, tenant_id=tenant_id)
    progress_dict = progress.model_dump()
    progress_dict['last_activity'] = progress_dict['last_activity'].isoformat()
    await db.user_progress.insert_one(progress_dict)
    
    # Log initial subscription history (starting as free)
    await db.subscription_history.insert_one({
        "user_id": user.id,
        "tenant_id": tenant_id,
        "action": "registration",
        "from_plan": None,
        "to_plan": "free",
        "timestamp": datetime.utcnow().isoformat()
    })
    
    # Send verification email
    email_sent = send_verification_email(
        to_email=user.email,
        user_name=user.full_name,
        verification_token=verification_token
    )
    
    if not email_sent:
        logger.warning(f"Failed to send verification email to {user.email}, auto-verifying")
        # Fallback: auto-verify if email fails
        await db.users.update_one(
            {"id": user.id},
            {"$set": {"email_verified": True}}
        )
    
    logger.info(f"New user registered: {user.email} for tenant: {tenant_id}")
    return user

@api_router.post("/auth/login", response_model=Token)
async def login(
    credentials: UserLogin,
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Login with tenant awareness.
    
    Users can only login on the tenant they registered with.
    """
    # Find user within this tenant
    user = await db.users.find_one({
        "email": credentials.email,
        "tenant_id": tenant_id
    })
    
    if not user:
        # Also check if user exists in another tenant (for better error message)
        other_tenant_user = await db.users.find_one({"email": credentials.email})
        if other_tenant_user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Account not found on this platform. Please login on the correct platform."
            )
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password"
        )
    
    # Verify password
    if not verify_password(credentials.password, user['hashed_password']):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password"
        )
    
    # Create tokens
    access_token = create_access_token(user['id'])
    refresh_token = create_refresh_token(user['id'])
    
    return Token(
        access_token=access_token,
        refresh_token=refresh_token
    )


@api_router.post("/auth/setup-token", response_model=APITokenResponse)
async def setup_token(
    credentials: SetupTokenRequest,
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Generate a long-lived API token for CLI/programmatic access.

    This endpoint authenticates the user and creates a persistent API token
    that can be used for non-interactive authentication (CI/CD, scripts, etc.).

    The full token is only returned once at creation time - store it securely.
    """
    # Find user within this tenant
    user = await db.users.find_one({
        "email": credentials.email,
        "tenant_id": tenant_id
    })

    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password"
        )

    # Verify password
    if not verify_password(credentials.password, user['hashed_password']):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password"
        )

    # Create the API token
    full_token, api_token = create_api_token_model(
        user_id=user['id'],
        name=credentials.token_name,
        scopes=[APITokenScope.READ, APITokenScope.WRITE],
        expires_days=365
    )

    # Store the token in database
    token_doc = api_token.model_dump()
    token_doc['tenant_id'] = tenant_id
    await db.api_tokens.insert_one(token_doc)

    logger.info(f"API token created for user {user['id']} (prefix: {api_token.token_prefix})")

    return APITokenResponse(
        token=full_token,
        token_prefix=api_token.token_prefix,
        name=api_token.name,
        scopes=api_token.scopes,
        expires_at=api_token.expires_at,
        created_at=api_token.created_at
    )


@api_router.get("/auth/tokens")
async def list_api_tokens(
    user_id: str = Depends(get_current_user)
):
    """List all API tokens for the current user."""
    tokens = await db.api_tokens.find({
        "user_id": user_id,
        "is_active": True
    }).to_list(100)

    return [
        {
            "id": token['id'],
            "name": token['name'],
            "token_prefix": token['token_prefix'],
            "scopes": token['scopes'],
            "created_at": token['created_at'],
            "last_used_at": token.get('last_used_at'),
            "expires_at": token.get('expires_at')
        }
        for token in tokens
    ]


@api_router.delete("/auth/tokens/{token_id}")
async def revoke_api_token(
    token_id: str,
    user_id: str = Depends(get_current_user)
):
    """Revoke an API token."""
    result = await db.api_tokens.update_one(
        {"id": token_id, "user_id": user_id},
        {"$set": {"is_active": False}}
    )

    if result.modified_count == 0:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Token not found"
        )

    logger.info(f"API token {token_id} revoked by user {user_id}")
    return {"status": "revoked"}


@api_router.post("/auth/forgot-password")
async def forgot_password(
    data: dict,
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Request password reset with tenant awareness."""
    import secrets
    
    email = data.get('email')
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")
    
    # Find user within this tenant (don't reveal if email exists for security)
    user = await db.users.find_one({
        "email": email,
        "tenant_id": tenant_id
    })
    
    if user:
        # Generate reset token
        reset_token = secrets.token_urlsafe(32)
        expires_at = datetime.utcnow() + timedelta(hours=1)
        
        # Store reset token with tenant_id
        await db.password_resets.delete_many({"email": email, "tenant_id": tenant_id})
        await db.password_resets.insert_one({
            "email": email,
            "tenant_id": tenant_id,
            "token": reset_token,
            "expires_at": expires_at.isoformat(),
            "created_at": datetime.utcnow().isoformat()
        })
        
        # Send password reset email
        email_sent = send_password_reset_email(
            to_email=email,
            user_name=user.get('full_name', 'User'),
            reset_token=reset_token
        )
        
        if email_sent:
            logger.info(f"Password reset email sent to {email} (tenant: {tenant_id})")
        else:
            logger.warning(f"Failed to send password reset email to {email}")
    
    # Always return success (don't reveal if email exists)
    return {"message": "If an account exists with this email, you will receive password reset instructions shortly."}

@api_router.post("/auth/reset-password")
async def reset_password(data: dict):
    """Reset password using token from email"""
    token = data.get('token')
    new_password = data.get('new_password')
    
    if not token or not new_password:
        raise HTTPException(status_code=400, detail="Token and new password are required")
    
    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    
    # Find reset token
    reset_record = await db.password_resets.find_one({"token": token})
    
    if not reset_record:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Check if expired
    expires_at = datetime.fromisoformat(reset_record['expires_at'])
    if datetime.utcnow() > expires_at:
        await db.password_resets.delete_one({"token": token})
        raise HTTPException(status_code=400, detail="Reset token has expired")
    
    # Update password
    hashed_password = hash_password(new_password)
    await db.users.update_one(
        {"email": reset_record['email']},
        {"$set": {"hashed_password": hashed_password}}
    )
    
    # Delete used token
    await db.password_resets.delete_one({"token": token})
    
    return {"message": "Password has been reset successfully. You can now login with your new password."}

@api_router.get("/auth/verify-email")
async def verify_email(token: str):
    """Verify email address using token from email link"""
    if not token:
        raise HTTPException(status_code=400, detail="Verification token is required")
    
    # Find user with this verification token
    user = await db.users.find_one({"verification_token": token})
    
    if not user:
        raise HTTPException(status_code=400, detail="Invalid or expired verification token")
    
    if user.get('email_verified'):
        return {"message": "Email already verified. You can login now."}
    
    # Update user as verified
    await db.users.update_one(
        {"verification_token": token},
        {
            "$set": {"email_verified": True},
            "$unset": {"verification_token": ""}
        }
    )
    
    logger.info(f"Email verified for user: {user.get('email')}")
    return {"message": "Email verified successfully! You can now login."}

@api_router.post("/auth/resend-verification")
async def resend_verification_email(
    data: dict,
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Resend email verification link"""
    email = data.get('email')
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")
    
    # Find user
    user = await db.users.find_one({
        "email": email,
        "tenant_id": tenant_id
    })
    
    if not user:
        # Don't reveal if email exists
        return {"message": "If an account exists with this email, a verification link will be sent."}
    
    if user.get('email_verified'):
        return {"message": "Email is already verified. You can login now."}
    
    # Generate new verification token
    new_token = str(uuid.uuid4())
    await db.users.update_one(
        {"email": email, "tenant_id": tenant_id},
        {"$set": {"verification_token": new_token}}
    )
    
    # Send verification email
    email_sent = send_verification_email(
        to_email=email,
        user_name=user.get('full_name', 'User'),
        verification_token=new_token
    )
    
    if email_sent:
        logger.info(f"Verification email resent to {email}")
    else:
        logger.warning(f"Failed to resend verification email to {email}")
    
    return {"message": "If an account exists with this email, a verification link will be sent."}

@api_router.get("/auth/me", response_model=User)
async def get_current_user_info(user_id: str = Depends(get_current_user)):
    """Get current user information"""
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "hashed_password": 0})
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    # Convert datetime strings
    if isinstance(user.get('created_at'), str):
        user['created_at'] = datetime.fromisoformat(user['created_at'])
    
    return User(**user)

# ============================================================================
# HELPER FUNCTIONS FOR PRIORITY QUESTION SYSTEM
# ============================================================================

import random

def randomize_answer_options(question_dict):
    """
    Randomize the order of answer options and return the new correct answer index.
    Returns: (randomized_options, new_correct_index, original_to_new_mapping)
    """
    options = question_dict.get('options', [])
    correct_answer = question_dict.get('correct_answer', 0)
    
    # Create list of (original_index, option_text)
    indexed_options = list(enumerate(options))
    
    # Shuffle the options
    random.shuffle(indexed_options)
    
    # Extract the new order
    new_options = [opt for _, opt in indexed_options]
    original_indices = [idx for idx, _ in indexed_options]
    
    # Find new correct answer index
    new_correct_index = original_indices.index(correct_answer)
    
    return new_options, new_correct_index, original_indices

async def get_user_unlock_status(user_id: str):
    """Check if user has unlocked the full question bank."""
    progress = await db.user_progress.find_one({"user_id": user_id}, {"_id": 0})
    if not progress:
        return False, 0
    return progress.get('full_bank_unlocked', False), progress.get('qualifying_sessions_completed', 0)

# Daily question limits by subscription tier
DAILY_LIMITS_BY_TIER = {
    'free': 50,
    'weekly': 200,
    'monthly': 500,
    'quarterly': -1,  # Unlimited
    'annual': -1,     # Unlimited
}

async def check_daily_question_limit(user_id: str) -> tuple:
    """
    Check if user has exceeded their daily question limit based on subscription tier.
    Returns: (can_continue, questions_remaining, is_subscriber, daily_limit)
    
    Tier limits:
    - Free: 50/day
    - Weekly: 200/day
    - Monthly: 500/day
    - Quarterly: Unlimited
    - Annual: Unlimited
    """
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        return False, 0, False, 50
    
    # Check subscription status and plan
    # The system uses subscription_plan for the tier name (weekly, monthly, etc.)
    # and subscription_status for active/free_grant/free
    subscription_status = user.get('subscription_status', 'free')
    subscription_plan = user.get('subscription_plan', 'free') or 'free'
    
    # User is a subscriber if they have an active subscription or free grant
    is_subscriber = subscription_status in ['active', 'free_grant'] and subscription_plan != 'free'
    
    # Determine the effective tier for limits
    if is_subscriber:
        effective_tier = subscription_plan
    else:
        effective_tier = 'free'
    
    # Get daily limit for this tier
    daily_limit = DAILY_LIMITS_BY_TIER.get(effective_tier, 50)
    
    # Unlimited tiers (quarterly and annual)
    if daily_limit == -1:
        return True, -1, True, -1
    
    today = datetime.utcnow().date().isoformat()
    
    # Get or create daily usage record
    daily_usage = await db.daily_usage.find_one({
        "user_id": user_id,
        "date": today
    }, {"_id": 0})
    
    if not daily_usage:
        # First question of the day
        await db.daily_usage.insert_one({
            "user_id": user_id,
            "date": today,
            "questions_viewed": 0,
            "created_at": datetime.utcnow().isoformat()
        })
        return True, daily_limit, is_subscriber, daily_limit
    
    questions_viewed = daily_usage.get('questions_viewed', 0)
    questions_remaining = max(0, daily_limit - questions_viewed)
    
    return questions_remaining > 0, questions_remaining, is_subscriber, daily_limit

async def increment_daily_usage(user_id: str, count: int = 1):
    """Increment the daily question count for a user."""
    today = datetime.utcnow().date().isoformat()
    await db.daily_usage.update_one(
        {"user_id": user_id, "date": today},
        {
            "$inc": {"questions_viewed": count},
            "$setOnInsert": {
                "user_id": user_id,
                "date": today,
                "created_at": datetime.utcnow().isoformat()
            }
        },
        upsert=True
    )

async def get_user_study_year(user_id: str) -> int:
    """Get user's current study year. Default to 2 if not set."""
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        return 2
    return user.get('current_year') or 2  # Default to year 2

async def get_user_category_progress(user_id: str, category: str) -> dict:
    """Get user's progress in a specific category to determine complexity level."""
    progress = await db.category_progress.find_one({
        "user_id": user_id,
        "category": category
    }, {"_id": 0})
    
    if not progress:
        return {"mastered_level": 0, "current_level": 1}  # Start at foundational
    
    return {
        "mastered_level": progress.get('mastered_level', 0),
        "current_level": progress.get('current_level', 1)
    }

# ============================================================================
# QUESTION ROUTES (Tenant-Aware)
# ============================================================================

@api_router.get("/questions", response_model=List[Question])
async def get_questions(
    request: Request,
    category: Optional[QuestionCategory] = None,
    difficulty: Optional[DifficultyLevel] = None,
    year: Optional[int] = None,
    source: Optional[str] = None,
    limit: int = 50,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Get questions with optional filters - randomized order with randomized answers.
    
    TENANT ISOLATION: Questions are filtered by tenant_id.
    
    FEATURES:
    - PRIORITY SYSTEM: Until user completes 3 qualifying sessions (85%+ on 50+ questions),
      only UNE priority questions are served.
    - YEAR FILTERING: Students only see questions up to their current study year
    - COMPLEXITY PROGRESSION: Start at foundational level, build up based on performance
    - DAILY LIMIT: Non-subscribers limited to 50 questions per day (checked when answering)
    """
    # NOTE: Daily limit is now checked when ANSWERING questions, not when fetching
    # This allows users to view questions but tracks actual usage
    
    # Check user's unlock status
    full_bank_unlocked, qualifying_sessions = await get_user_unlock_status(user_id)
    
    # Get user's study year for filtering
    user_study_year = await get_user_study_year(user_id)
    
    # Start building query with $and to properly combine tenant filter with other filters
    # Tenant filter - questions must belong to this tenant or be global
    tenant_filter = {
        '$or': [
            {'tenant_id': tenant_id},
            {'tenant_id': {'$exists': False}},  # Legacy questions without tenant_id
            {'tenant_id': None}  # Global questions
        ]
    }
    
    # Build additional filters
    additional_filters = []
    
    if category:
        additional_filters.append({'category': category.value})
        
        # Get user's progress in this category for complexity progression
        cat_progress = await get_user_category_progress(user_id, category.value)
        current_level = cat_progress.get('current_level', 1)
        
        # If no explicit difficulty requested, use progressive difficulty
        if not difficulty:
            # Allow questions at or below current level (with some stretch questions)
            additional_filters.append({'difficulty': {'$lte': str(min(current_level + 1, 4))}})
    
    if difficulty:
        additional_filters.append({'difficulty': difficulty.value})
    
    # YEAR FILTERING: Only show questions up to user's current study year
    # If user explicitly requests a specific year, use that (but cap at their level)
    if year:
        # User requested specific year - cap at their study year
        effective_year = min(year, user_study_year)
        additional_filters.append({'year': effective_year})
    else:
        # No year specified - show all years up to their study year
        additional_filters.append({'year': {'$lte': user_study_year}})
    
    # PRIORITY SYSTEM: If not unlocked, only serve UNE priority questions
    if not full_bank_unlocked:
        additional_filters.append({'source': 'une_priority'})
    else:
        # Source filtering for unlocked users
        if source and source != 'all':
            if source == 'imported':
                additional_filters.append({'user_id': user_id})
                additional_filters.append({'source': 'imported'})
            elif source == 'shared':
                additional_filters.append({'source': 'shared'})
            elif source == 'sample':
                additional_filters.append({'source': 'sample'})
            elif source == 'une_priority':
                additional_filters.append({'source': 'une_priority'})
        # If source is 'all' or not specified, include all sources (no filter added)
    
    # Exclude quarantined questions
    additional_filters.append({'quarantined': {'$ne': True}})
    
    # Combine tenant filter with additional filters using $and
    query = {
        '$and': [tenant_filter] + additional_filters
    }
    
    questions = await db.questions.find(query, {"_id": 0}).to_list(limit * 3)
    
    # Convert datetime strings
    for q in questions:
        if isinstance(q.get('created_at'), str):
            q['created_at'] = datetime.fromisoformat(q['created_at'])
    
    # Randomize question order
    random.shuffle(questions)
    
    # Limit results
    questions = questions[:limit]
    
    # NOTE: Daily usage is now tracked when answering questions, not when fetching
    # This prevents counting all fetched questions against the limit
    
    # Randomize answer options for each question
    randomized_questions = []
    for q in questions:
        new_options, new_correct, _ = randomize_answer_options(q)
        q['options'] = new_options
        q['correct_answer'] = new_correct
        randomized_questions.append(Question(**q))
    
    return randomized_questions

@api_router.get("/questions/daily-limit")
async def get_daily_limit_status(user_id: str = Depends(get_current_user)):
    """Get user's daily question limit status."""
    can_continue, questions_remaining, is_subscriber, daily_limit = await check_daily_question_limit(user_id)
    
    return {
        "can_continue": can_continue,
        "questions_remaining": questions_remaining,
        "is_subscriber": is_subscriber,
        "daily_limit": daily_limit
    }

@api_router.get("/questions/adaptive", response_model=List[Question])
async def get_adaptive_questions(
    request: Request,
    category: Optional[QuestionCategory] = None,
    count: int = 10,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Get questions adapted to user's current level with randomized answers.
    
    TENANT ISOLATION: Questions are filtered by tenant_id.
    PRIORITY SYSTEM: Until user completes 3 qualifying sessions, only UNE questions.
    """
    # Check user's unlock status
    full_bank_unlocked, _ = await get_user_unlock_status(user_id)
    
    # Get user progress
    progress = await db.user_progress.find_one({"user_id": user_id}, {"_id": 0})
    if not progress:
        progress = UserProgress(user_id=user_id, tenant_id=tenant_id).model_dump()
    else:
        if isinstance(progress.get('last_activity'), str):
            progress['last_activity'] = datetime.fromisoformat(progress['last_activity'])
    
    progress_obj = UserProgress(**progress)
    
    # Build query with tenant filter
    tenant_filter = {
        '$or': [
            {'tenant_id': tenant_id},
            {'tenant_id': {'$exists': False}},
            {'tenant_id': None}
        ]
    }
    
    if not full_bank_unlocked:
        query = {'source': 'une_priority', **tenant_filter}
    else:
        query = tenant_filter
    
    if category:
        query['category'] = category.value
        
    all_questions = await db.questions.find(query, {"_id": 0}).to_list(1000)
    
    # Convert datetime
    for q in all_questions:
        if isinstance(q.get('created_at'), str):
            q['created_at'] = datetime.fromisoformat(q['created_at'])
    
    questions_obj = [Question(**q) for q in all_questions]
    
    # If no questions from adaptive engine, return random questions
    if len(questions_obj) == 0:
        return []
    
    # Use adaptive engine to select questions
    selected = adaptive_engine.get_next_questions(
        progress_obj,
        questions_obj,
        category,
        count
    )
    
    # If adaptive selection returns nothing (new user), return random easy questions
    if len(selected) == 0:
        # Filter for easy questions or just shuffle all
        easy_questions = [q for q in questions_obj if q.difficulty == DifficultyLevel.EASY]
        if len(easy_questions) > 0:
            random.shuffle(easy_questions)
            selected = easy_questions[:count]
        else:
            # No easy questions, just return random selection
            random.shuffle(questions_obj)
            selected = questions_obj[:count]
    
    # Randomize answer options for each question
    randomized_questions = []
    for q in selected:
        q_dict = q.model_dump()
        new_options, new_correct, _ = randomize_answer_options(q_dict)
        q_dict['options'] = new_options
        q_dict['correct_answer'] = new_correct
        randomized_questions.append(Question(**q_dict))
    
    return randomized_questions

@api_router.post("/questions/answer")
async def submit_answer(
    request: Request,
    answer: UserAnswer,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Submit an answer and update progress with tenant awareness."""
    # Check daily limit based on subscription tier BEFORE processing the answer
    can_continue, questions_remaining, is_subscriber, daily_limit = await check_daily_question_limit(user_id)
    
    if not can_continue:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Daily question limit reached. Upgrade your subscription for more questions!"
        )
    
    # Increment daily usage for users with limits (not quarterly/annual)
    if daily_limit != -1:
        await increment_daily_usage(user_id, 1)
    
    # Get question
    question = await db.questions.find_one({"id": answer.question_id}, {"_id": 0})
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    if isinstance(question.get('created_at'), str):
        question['created_at'] = datetime.fromisoformat(question['created_at'])
    
    question_obj = Question(**question)
    
    # Get user progress
    progress = await db.user_progress.find_one({"user_id": user_id}, {"_id": 0})
    if not progress:
        progress = UserProgress(user_id=user_id).model_dump()
    else:
        if isinstance(progress.get('last_activity'), str):
            progress['last_activity'] = datetime.fromisoformat(progress['last_activity'])
    
    progress_obj = UserProgress(**progress)
    
    # Process answer with adaptive engine
    updated_progress = adaptive_engine.process_answer(progress_obj, answer, question_obj)
    
    # Update database
    progress_dict = updated_progress.model_dump()
    progress_dict['last_activity'] = progress_dict['last_activity'].isoformat()
    
    await db.user_progress.update_one(
        {"user_id": user_id},
        {"$set": progress_dict},
        upsert=True
    )
    
    # Record answer in history
    answer_dict = answer.model_dump()
    answer_dict['user_id'] = user_id
    answer_dict['timestamp'] = answer_dict['timestamp'].isoformat()
    await db.answer_history.insert_one(answer_dict)
    
    # Update today's study session
    today = datetime.utcnow().date().isoformat()
    await db.study_sessions.update_one(
        {"user_id": user_id, "date": today},
        {
            "$inc": {
                "questions_answered": 1,
                "correct_answers": 1 if answer.is_correct else 0,
                "time_spent": answer.time_taken
            },
            "$addToSet": {"categories_studied": question_obj.category.value}
        },
        upsert=True
    )
    
    # Calculate remaining questions (decrement by 1 if has limit)
    new_remaining = questions_remaining - 1 if daily_limit != -1 else -1
    
    return {
        "success": True,
        "current_difficulty": updated_progress.category_progress.get(
            question_obj.category.value,
            CategoryProgress(category=question_obj.category)
        ).current_difficulty,
        "current_streak": updated_progress.current_streak,
        "questions_remaining": new_remaining,
        "daily_limit": daily_limit
    }

@api_router.post("/session/finish")
async def finish_session(user_id: str = Depends(get_current_user)):
    """Finish current study session and return session summary"""
    # Get today's session
    today = datetime.utcnow().date().isoformat()
    session = await db.study_sessions.find_one(
        {"user_id": user_id, "date": today},
        {"_id": 0}
    )
    
    if not session:
        return {
            "success": True,
            "message": "No active session",
            "session_data": None
        }
    
    # Get user progress for additional stats
    progress = await db.user_progress.find_one({"user_id": user_id}, {"_id": 0})
    
    session_summary = {
        "date": session.get("date"),
        "questions_answered": session.get("questions_answered", 0),
        "correct_answers": session.get("correct_answers", 0),
        "accuracy": round((session.get("correct_answers", 0) / session.get("questions_answered", 1)) * 100, 1) if session.get("questions_answered", 0) > 0 else 0,
        "time_spent": session.get("time_spent", 0),
        "categories_studied": session.get("categories_studied", []),
        "current_streak": progress.get("current_streak", 0) if progress else 0,
        "total_questions_all_time": progress.get("total_questions_answered", 0) if progress else 0
    }
    
    return {
        "success": True,
        "message": "Session finished successfully",
        "session_data": session_summary
    }

@api_router.post("/questions/import")
async def import_questions(
    request: Request,
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Import questions from Excel/CSV/Word/PDF file with duplicate detection and tenant assignment."""
    try:
        # Read file
        contents = await file.read()
        filename_lower = file.filename.lower()

        # Parse based on file type
        if filename_lower.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(contents))
        elif filename_lower.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(io.BytesIO(contents))
        elif filename_lower.endswith(('.docx', '.doc')):
            # For Word documents, try to extract text and parse
            try:
                from docx import Document
                doc = Document(io.BytesIO(contents))

                # Try to find tables in the document
                if doc.tables:
                    # Extract data from the first table
                    table = doc.tables[0]
                    data = []
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    for row in table.rows[1:]:
                        row_data = [cell.text.strip() for cell in row.cells]
                        data.append(row_data)
                    df = pd.DataFrame(data, columns=headers)
                else:
                    raise HTTPException(
                        status_code=400,
                        detail="Word document must contain a table with question data. Please use Excel/CSV format or ensure your Word document has a properly formatted table."
                    )
            except ImportError:
                raise HTTPException(
                    status_code=400,
                    detail="Word document support not available. Please install python-docx package or use Excel/CSV format instead."
                )
        elif filename_lower.endswith('.pdf'):
            # For PDF files, try to extract tables
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(contents)) as pdf:
                    all_tables = []
                    for page in pdf.pages:
                        tables = page.extract_tables()
                        if tables:
                            all_tables.extend(tables)

                    if not all_tables:
                        raise HTTPException(
                            status_code=400,
                            detail="No tables found in PDF. Please ensure your PDF contains a properly formatted table with question data, or use Excel/CSV format instead."
                        )

                    # Use the first table
                    table = all_tables[0]
                    headers = table[0]
                    data = table[1:]
                    df = pd.DataFrame(data, columns=headers)
            except ImportError:
                raise HTTPException(
                    status_code=400,
                    detail="PDF support not available. Please install pdfplumber package or use Excel/CSV format instead."
                )
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format. Supported formats: Excel (.xlsx, .xls), CSV (.csv), Word (.docx, .doc), PDF (.pdf)"
            )
        
        # Validate columns - updated with new format including year
        required_columns = [
            'question', 'optionA', 'optionB', 'optionC', 'optionD', 'optionE',
            'correctAnswer', 'explanation', 'category', 'subCategory', 'year', 'level'
        ]
        
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise HTTPException(
                status_code=400,
                detail=f"Missing columns: {', '.join(missing_columns)}. Required columns are: {', '.join(required_columns)}"
            )
        
        # Get all existing question texts for this tenant for duplicate detection
        existing_questions = await db.questions.find(
            {"tenant_id": tenant_id},
            {"question": 1, "_id": 0}
        ).to_list(100000)
        existing_question_texts = set(q['question'].strip().lower() for q in existing_questions)
        
        # Track duplicates
        duplicates_skipped = 0
        duplicates_in_file = 0
        questions_in_file = set()  # Track questions within the file to avoid duplicates in same upload
        
        # Create questions
        questions = []
        for _, row in df.iterrows():
            try:
                question_text = str(row['question']).strip()
                question_text_lower = question_text.lower()
                
                # Check for duplicate in database
                if question_text_lower in existing_question_texts:
                    duplicates_skipped += 1
                    logger.info(f"Skipping duplicate question (already in database): {question_text[:50]}...")
                    continue
                
                # Check for duplicate within the same file
                if question_text_lower in questions_in_file:
                    duplicates_in_file += 1
                    logger.info(f"Skipping duplicate question (duplicate in file): {question_text[:50]}...")
                    continue
                
                questions_in_file.add(question_text_lower)
                
                # Parse options - always include all 5 options
                options = [
                    str(row['optionA']),
                    str(row['optionB']),
                    str(row['optionC']),
                    str(row['optionD']),
                    str(row['optionE'])
                ]
                
                # Parse correct answer (A, B, C, D, or E)
                correct_letter = str(row['correctAnswer']).upper().strip()
                correct_index = ord(correct_letter) - ord('A')
                
                # Get sub-category
                sub_category = str(row['subCategory']) if pd.notna(row['subCategory']) else None
                
                # Get year (1-6)
                year_val = int(row['year']) if pd.notna(row['year']) else 2
                year_val = max(1, min(6, year_val))  # Clamp to 1-6
                
                # Create question with tenant_id
                question = Question(
                    question=question_text,
                    options=options,
                    correct_answer=correct_index,
                    explanation=str(row['explanation']),
                    category=QuestionCategory(str(row['category']).lower()),
                    sub_category=sub_category,
                    year=year_val,
                    difficulty=DifficultyLevel(str(row['level'])),
                    user_id=user_id,
                    source="imported",
                    tenant_id=tenant_id  # Multi-tenant: assign to current tenant
                )
                
                questions.append(question)
            except Exception as e:
                logger.warning(f"Skipping row due to error: {e}")
                continue
        
        # Store in database
        if questions:
            questions_dict = [q.model_dump() for q in questions]
            for q in questions_dict:
                q['created_at'] = q['created_at'].isoformat()
            
            await db.questions.insert_many(questions_dict)
        
        # Build response message
        total_in_file = len(df)
        message_parts = [f"Successfully imported {len(questions)} questions"]
        if duplicates_skipped > 0:
            message_parts.append(f"{duplicates_skipped} duplicates skipped (already in database)")
        if duplicates_in_file > 0:
            message_parts.append(f"{duplicates_in_file} duplicates within file skipped")
        
        logger.info(f"Imported {len(questions)} questions for tenant {tenant_id}")
        return {
            "success": True,
            "questions_imported": len(questions),
            "duplicates_skipped": duplicates_skipped,
            "duplicates_in_file": duplicates_in_file,
            "total_in_file": total_in_file,
            "message": ". ".join(message_parts)
        }
    
    except Exception as e:
        logger.error(f"Error importing questions: {e}")
        raise HTTPException(status_code=400, detail=str(e))

@api_router.post("/questions/report")
async def report_question(
    request: Request,
    report: QuestionReport,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Report a question for review with tenant awareness."""
    report.user_id = user_id
    report.tenant_id = tenant_id  # Multi-tenant: assign to current tenant
    
    report_dict = report.model_dump()
    report_dict['timestamp'] = report_dict['timestamp'].isoformat()
    report_dict['status'] = 'pending'  # Add status field
    
    await db.question_reports.insert_one(report_dict)
    
    # Get question details and user info for email
    question = await db.questions.find_one({"id": report.question_id}, {"_id": 0})
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    
    # Send email notification to support
    if question and user:
        send_question_report_notification(
            question_id=report.question_id,
            question_text=question.get('question', 'N/A'),
            report_reason=report.reason,
            reporter_email=user.get('email', 'Unknown'),
            reporter_name=user.get('full_name', 'Unknown User')
        )
    
    # Check if this question has 4+ reports within this tenant
    report_count = await db.question_reports.count_documents({
        "question_id": report.question_id,
        "tenant_id": tenant_id,
        "resolved": False
    })
    
    # Quarantine question if 4+ reports
    if report_count >= 4:
        await db.questions.update_one(
            {"id": report.question_id},
            {"$set": {"quarantined": True}}
        )
        # Update report status
        await db.question_reports.update_many(
            {"question_id": report.question_id, "tenant_id": tenant_id},
            {"$set": {"status": "quarantined"}}
        )
        return {
            "success": True, 
            "message": "Report submitted. Question has been quarantined for review.",
            "quarantined": True
        }
    
    return {
        "success": True, 
        "message": f"Report submitted successfully. ({report_count}/4 reports)",
        "quarantined": False
    }

# ============================================================================
# AI GENERATION ROUTES
# ============================================================================

@api_router.post("/ai/generate", response_model=AIGenerationResponse)
async def generate_ai_questions(
    request: AIGenerationRequest,
    user_id: str = Depends(get_current_user)
):
    """Generate questions using AI"""
    # Get user
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Check daily usage
    if user['ai_daily_uses'] >= user['ai_max_daily_uses']:
        raise HTTPException(
            status_code=429,
            detail="Daily AI usage limit reached"
        )
    
    # Generate questions
    questions = await ai_service.generate_questions(request, user_id)
    
    if not questions:
        raise HTTPException(
            status_code=500,
            detail="Failed to generate questions. Please try again."
        )
    
    # Store questions
    questions_dict = [q.model_dump() for q in questions]
    for q in questions_dict:
        q['created_at'] = q['created_at'].isoformat()
    
    await db.questions.insert_many(questions_dict)
    
    # Update user's AI usage
    await db.users.update_one(
        {"id": user_id},
        {"$inc": {"ai_daily_uses": 1}}
    )
    
    # Get updated usage
    updated_user = await db.users.find_one({"id": user_id}, {"_id": 0})
    
    return AIGenerationResponse(
        questions=questions,
        usage_count=updated_user['ai_daily_uses'],
        remaining_daily_uses=updated_user['ai_max_daily_uses'] - updated_user['ai_daily_uses']
    )

@api_router.get("/ai/usage")
async def get_ai_usage(user_id: str = Depends(get_current_user)):
    """Get AI usage statistics"""
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {
        "daily_uses": user['ai_daily_uses'],
        "max_daily_uses": user['ai_max_daily_uses'],
        "remaining_uses": user['ai_max_daily_uses'] - user['ai_daily_uses'],
        "storage_used_gb": user['storage_used_gb'],
        "storage_quota_gb": user['storage_quota_gb']
    }

# ============================================================================
# ANALYTICS ROUTES (Tenant-Aware)
# ============================================================================

@api_router.get("/analytics", response_model=UserAnalytics)
async def get_analytics(
    request: Request,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Get user analytics with tenant awareness."""
    # Get progress for this user (user-scoped, tenant inherited from user)
    progress = await db.user_progress.find_one({"user_id": user_id}, {"_id": 0})
    if not progress:
        raise HTTPException(status_code=404, detail="Progress not found")
    
    # Get study sessions for this user
    sessions = await db.study_sessions.find({"user_id": user_id}, {"_id": 0}).to_list(1000)
    
    # Calculate statistics
    total_questions = progress['total_questions_answered']
    correct_rate = (progress['total_correct'] / total_questions * 100) if total_questions > 0 else 0
    avg_time = progress['total_time_spent'] / total_questions if total_questions > 0 else 0
    study_days = len(sessions)
    
    # Category performance
    category_performance = {}
    for cat_name, cat_data in progress.get('category_progress', {}).items():
        total = cat_data['total_answered']
        correct = cat_data['total_correct']
        accuracy = (correct / total * 100) if total > 0 else 0
        
        category_performance[cat_name] = {
            "total_answered": total,
            "correct": correct,
            "accuracy": accuracy,
            "current_difficulty": cat_data['current_difficulty']
        }
    
    return UserAnalytics(
        total_questions=total_questions,
        correct_rate=correct_rate,
        average_time_per_question=avg_time,
        study_days=study_days,
        current_streak=progress['current_streak'],
        highest_streak=progress['highest_streak'],
        category_performance=category_performance
    )

@api_router.get("/analytics/sessions")
async def get_study_sessions(
    request: Request,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Get study sessions with tenant awareness."""
    # Sessions are user-scoped
    query = {"user_id": user_id}
    
    if start_date:
        query['date'] = {"$gte": start_date}
    if end_date:
        if 'date' in query:
            query['date']['$lte'] = end_date
        else:
            query['date'] = {"$lte": end_date}
    
    sessions = await db.study_sessions.find(query, {"_id": 0}).to_list(1000)
    return sessions

@api_router.get("/analytics/subcategory")
async def get_subcategory_analytics(
    request: Request,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Get sub-category performance analytics with tenant awareness."""
    # Get answer history for this user (user-scoped)
    answers = await db.answer_history.find({"user_id": user_id}, {"_id": 0}).to_list(10000)
    
    # Get question details for sub-categories - filtered by tenant
    question_ids = [a.get('question_id') for a in answers if a.get('question_id')]
    questions = await db.questions.find({
        "id": {"$in": question_ids},
        "$or": [
            {"tenant_id": tenant_id},
            {"tenant_id": {"$exists": False}},
            {"tenant_id": None}
        ]
    }, {"_id": 0}).to_list(10000)
    
    # Create question lookup
    question_lookup = {q['id']: q for q in questions}
    
    # Calculate sub-category stats
    subcategory_stats = {}
    for answer in answers:
        q_id = answer.get('question_id')
        if q_id and q_id in question_lookup:
            question = question_lookup[q_id]
            category = question.get('category', 'general')
            sub_cat = question.get('sub_category') or 'General'
            key = f"{category}|{sub_cat}"
            
            if key not in subcategory_stats:
                subcategory_stats[key] = {
                    'category': category,
                    'sub_category': sub_cat,
                    'total_answered': 0,
                    'total_correct': 0,
                    'accuracy': 0
                }
            
            subcategory_stats[key]['total_answered'] += 1
            if answer.get('is_correct'):
                subcategory_stats[key]['total_correct'] += 1
    
    # Calculate accuracy for each sub-category
    for key in subcategory_stats:
        stats = subcategory_stats[key]
        if stats['total_answered'] > 0:
            stats['accuracy'] = round((stats['total_correct'] / stats['total_answered']) * 100, 1)
    
    # Convert to list and sort by category then sub_category
    result = list(subcategory_stats.values())
    result.sort(key=lambda x: (x['category'], x['sub_category']))
    
    return result

# ============================================================================
# EXAM MODE ROUTES (Tenant-Aware)
# ============================================================================

@api_router.post("/exam/start", response_model=ExamSession)
async def start_exam(
    request: Request,
    question_count: int = 50,
    time_limit: int = 60,
    category: Optional[QuestionCategory] = None,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Start an exam session with randomized questions and answers.
    
    TENANT ISOLATION: Only questions from current tenant are used.
    PRIORITY SYSTEM: Until unlocked, only UNE questions are available.
    Minimum 50 questions required for qualifying sessions.
    """
    # Check user's unlock status
    full_bank_unlocked, _ = await get_user_unlock_status(user_id)
    
    # Build query with tenant filter
    tenant_filter = {
        '$or': [
            {'tenant_id': tenant_id},
            {'tenant_id': {'$exists': False}},
            {'tenant_id': None}
        ]
    }
    
    if not full_bank_unlocked:
        query = {'source': 'une_priority', 'quarantined': {'$ne': True}, **tenant_filter}
    else:
        query = {'quarantined': {'$ne': True}, **tenant_filter}
    
    if category:
        query['category'] = category.value
    
    questions = await db.questions.find(query, {"_id": 0}).to_list(5000)
    
    # Shuffle questions
    random.shuffle(questions)
    selected = questions[:question_count]
    
    # Randomize answer options and store mappings
    answer_mappings = {}
    for q in selected:
        new_options, new_correct, original_indices = randomize_answer_options(q)
        q['options'] = new_options
        q['correct_answer'] = new_correct
        answer_mappings[q['id']] = original_indices
    
    # Create exam session with tenant_id
    exam = ExamSession(
        user_id=user_id,
        tenant_id=tenant_id,  # Multi-tenant: assign to current tenant
        question_ids=[q['id'] for q in selected],
        answers={},
        time_limit=time_limit,
        answer_mappings=answer_mappings,
        question_count=len(selected)
    )
    
    exam_dict = exam.model_dump()
    exam_dict['started_at'] = exam_dict['started_at'].isoformat()
    
    await db.exam_sessions.insert_one(exam_dict)
    
    return exam

@api_router.get("/exam/{exam_id}/questions")
async def get_exam_questions(
    request: Request,
    exam_id: str,
    user_id: str = Depends(get_current_user),
    tenant_id: str = Depends(get_current_tenant)
):
    """Get questions for an exam session with randomized answers."""
    # Verify exam belongs to user and tenant
    exam = await db.exam_sessions.find_one({
        "id": exam_id,
        "user_id": user_id,
        "$or": [
            {"tenant_id": tenant_id},
            {"tenant_id": {"$exists": False}}
        ]
    }, {"_id": 0})
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    # Get questions
    questions = await db.questions.find(
        {"id": {"$in": exam['question_ids']}},
        {"_id": 0}
    ).to_list(1000)
    
    # Apply the same answer randomization stored in the exam session
    answer_mappings = exam.get('answer_mappings', {})
    
    randomized_questions = []
    for q in questions:
        if isinstance(q.get('created_at'), str):
            q['created_at'] = datetime.fromisoformat(q['created_at'])
        
        # Apply stored randomization or create new one
        if q['id'] in answer_mappings:
            original_indices = answer_mappings[q['id']]
            original_options = q['options']
            q['options'] = [original_options[i] for i in original_indices]
            q['correct_answer'] = original_indices.index(q['correct_answer'])
        
        randomized_questions.append(q)
    
    # Sort by the order in question_ids
    question_order = {qid: idx for idx, qid in enumerate(exam['question_ids'])}
    randomized_questions.sort(key=lambda x: question_order.get(x['id'], 999))
    
    return randomized_questions

@api_router.post("/exam/{exam_id}/answer")
async def submit_exam_answer(
    exam_id: str,
    question_id: str,
    answer: int,
    user_id: str = Depends(get_current_user)
):
    """Submit answer for an exam question"""
    await db.exam_sessions.update_one(
        {"id": exam_id, "user_id": user_id},
        {"$set": {f"answers.{question_id}": answer}}
    )
    
    return {"success": True}

@api_router.post("/exam/{exam_id}/complete", response_model=ExamResult)
async def complete_exam(
    exam_id: str,
    user_id: str = Depends(get_current_user)
):
    """Complete exam and get results. Check for qualifying session."""
    exam = await db.exam_sessions.find_one({"id": exam_id, "user_id": user_id}, {"_id": 0})
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    # Get questions
    questions = await db.questions.find(
        {"id": {"$in": exam['question_ids']}},
        {"_id": 0}
    ).to_list(1000)
    
    # Get answer mappings for correct answer calculation
    answer_mappings = exam.get('answer_mappings', {})
    
    # Calculate score
    correct = 0
    detailed_results = []
    
    for q in questions:
        user_answer = exam['answers'].get(q['id'])
        
        # Determine correct answer based on randomization
        if q['id'] in answer_mappings:
            original_indices = answer_mappings[q['id']]
            correct_answer_in_randomized = original_indices.index(q['correct_answer'])
        else:
            correct_answer_in_randomized = q['correct_answer']
        
        is_correct = user_answer == correct_answer_in_randomized
        if is_correct:
            correct += 1
        
        detailed_results.append({
            "question_id": q['id'],
            "question": q['question'],
            "user_answer": user_answer,
            "correct_answer": correct_answer_in_randomized,
            "is_correct": is_correct,
            "explanation": q['explanation']
        })
    
    total = len(questions)
    score = int((correct / total) * 100) if total > 0 else 0
    
    # Calculate time taken
    started_at = datetime.fromisoformat(exam['started_at'])
    time_taken = int((datetime.utcnow() - started_at).total_seconds())
    
    # Check for qualifying session (85%+ on 50+ questions)
    is_qualifying = False
    qualifying_message = None
    
    if total >= 50 and score >= 85:
        is_qualifying = True
        
        # Update user progress - increment qualifying sessions
        progress = await db.user_progress.find_one({"user_id": user_id}, {"_id": 0})
        current_qualifying = progress.get('qualifying_sessions_completed', 0) if progress else 0
        new_qualifying = current_qualifying + 1
        
        update_data = {
            "qualifying_sessions_completed": new_qualifying
        }
        
        # Check if this unlocks the full bank (3 qualifying sessions)
        if new_qualifying >= 3:
            update_data['full_bank_unlocked'] = True
            qualifying_message = "🎉 Congratulations! You've completed 3 qualifying sessions and unlocked the FULL question bank!"
        else:
            remaining = 3 - new_qualifying
            qualifying_message = f"✅ Qualifying session completed! {remaining} more qualifying session(s) needed to unlock the full question bank."
        
        await db.user_progress.update_one(
            {"user_id": user_id},
            {"$set": update_data},
            upsert=True
        )
        
        # Send qualifying session email notification
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if user and user.get('email'):
            send_qualifying_session_email(
                to_email=user['email'],
                user_name=user.get('full_name', 'User'),
                sessions_completed=new_qualifying,
                score=score
            )
    
    # Update exam
    await db.exam_sessions.update_one(
        {"id": exam_id},
        {
            "$set": {
                "completed_at": datetime.utcnow().isoformat(),
                "score": score,
                "is_qualifying_session": is_qualifying
            }
        }
    )
    
    result = ExamResult(
        score=score,
        correct=correct,
        total=total,
        time_taken=time_taken,
        detailed_results=detailed_results
    )
    
    # Add qualifying info to response
    return {
        **result.model_dump(),
        "is_qualifying_session": is_qualifying,
        "qualifying_message": qualifying_message
    }

# ============================================================================
# DATA EXPORT ROUTES
# ============================================================================

@api_router.post("/export/data")
async def export_user_data(
    request: DataExportRequest,
    user_id: str = Depends(get_current_user)
):
    """Export all user data as ZIP file"""
    # Create ZIP file in memory
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Export questions
        if request.include_questions:
            questions = await db.questions.find({"user_id": user_id}, {"_id": 0}).to_list(10000)
            for q in questions:
                if isinstance(q.get('created_at'), str):
                    q['created_at'] = q['created_at']
            
            zip_file.writestr(
                "questions.json",
                json.dumps(questions, indent=2, default=str)
            )
        
        # Export progress
        if request.include_progress:
            progress = await db.user_progress.find_one({"user_id": user_id}, {"_id": 0})
            if progress:
                zip_file.writestr(
                    "progress.json",
                    json.dumps(progress, indent=2, default=str)
                )
        
        # Export sessions
        if request.include_sessions:
            sessions = await db.study_sessions.find({"user_id": user_id}, {"_id": 0}).to_list(10000)
            zip_file.writestr(
                "sessions.json",
                json.dumps(sessions, indent=2, default=str)
            )
        
        # Export answer history
        answers = await db.answer_history.find({"user_id": user_id}, {"_id": 0}).to_list(10000)
        zip_file.writestr(
            "answer_history.json",
            json.dumps(answers, indent=2, default=str)
        )
    
    # Prepare for download
    zip_buffer.seek(0)
    
    return StreamingResponse(
        io.BytesIO(zip_buffer.read()),
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename=medmcq_data_{user_id}.zip"
        }
    )

# ============================================================================
# SAMPLE DATA INITIALIZATION
# ============================================================================

@api_router.post("/init/sample-data")
async def initialize_sample_data(
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Initialize comprehensive sample questions (for demo purposes)"""
    # Clear existing sample data for this tenant
    await db.questions.delete_many({"source": "sample", "tenant_id": tenant_id})
    
    comprehensive_questions = [
        # Respiratory
        {"id": "resp-001", "question": "A 65-year-old man with a 40-pack-year smoking history presents with worsening dyspnea and a chronic cough. Pulmonary function tests show a decreased FEV1/FVC ratio, increased total lung capacity, and decreased diffusing capacity. Which of the following is the most likely diagnosis?", "options": ["Asthma", "Chronic bronchitis", "Emphysema", "Bronchiectasis", "Interstitial lung disease"], "correct_answer": 2, "explanation": "The patient's presentation and pulmonary function tests are classic for emphysema. Key findings include: 1) Decreased FEV1/FVC ratio indicating obstructive lung disease, 2) Increased total lung capacity indicating hyperinflation, and 3) Decreased diffusing capacity due to destruction of alveolar walls.", "category": "respiratory", "year": 2, "difficulty": "2", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        {"id": "resp-002", "question": "A 28-year-old woman presents to the emergency department with sudden onset of dyspnea and right-sided pleuritic chest pain. She recently returned from a long international flight. Her vital signs show tachycardia and mild hypoxemia. What is the most appropriate initial diagnostic test?", "options": ["Chest X-ray", "CT pulmonary angiography", "D-dimer", "Ventilation-perfusion scan", "Arterial blood gas"], "correct_answer": 2, "explanation": "Given the clinical suspicion for pulmonary embolism (recent travel, pleuritic chest pain, dyspnea), D-dimer is the most appropriate initial test. A negative D-dimer can rule out PE in low-risk patients. If positive, proceed to CT pulmonary angiography for definitive diagnosis.", "category": "respiratory", "year": 3, "difficulty": "2", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        # Cardiology
        {"id": "cardio-001", "question": "A 55-year-old woman presents with crushing substernal chest pain that radiates to her left arm. ECG shows ST-segment elevation in leads II, III, and aVF. Which coronary artery is most likely occluded?", "options": ["Left anterior descending artery", "Left circumflex artery", "Right coronary artery", "Left main coronary artery", "Posterior descending artery"], "correct_answer": 2, "explanation": "ST-segment elevation in leads II, III, and aVF indicates an inferior wall myocardial infarction. The right coronary artery (RCA) supplies the inferior wall of the left ventricle in 85% of people.", "category": "cardiology", "year": 3, "difficulty": "3", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        {"id": "cardio-002", "question": "A 45-year-old man with hypertension presents for a routine check-up. His blood pressure is 145/95 mmHg. Which of the following is the most appropriate first-line antihypertensive medication for this patient without other comorbidities?", "options": ["Thiazide diuretic", "Beta-blocker", "Calcium channel blocker", "Alpha-blocker", "ACE inhibitor"], "correct_answer": 0, "explanation": "Thiazide diuretics are recommended as first-line therapy for uncomplicated hypertension according to JNC guidelines. They have been shown to reduce cardiovascular events and are cost-effective.", "category": "cardiology", "year": 2, "difficulty": "1", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        # Neurology
        {"id": "neuro-001", "question": "A 45-year-old man with chronic alcohol use presents with confusion, ataxia, and nystagmus. Which vitamin deficiency is most likely responsible?", "options": ["Vitamin B1 (Thiamine)", "Vitamin B12", "Vitamin C", "Vitamin D", "Folate"], "correct_answer": 0, "explanation": "Wernicke's encephalopathy, characterized by the triad of confusion, ataxia, and nystagmus, is caused by thiamine (Vitamin B1) deficiency, commonly seen in chronic alcoholics.", "category": "neurology", "year": 2, "difficulty": "2", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        {"id": "neuro-002", "question": "A 72-year-old woman presents with sudden onset of right-sided weakness and inability to speak. CT scan of the head shows no hemorrhage. What is the time window for administering tissue plasminogen activator (tPA)?", "options": ["1.5 hours", "3 hours", "4.5 hours", "6 hours", "12 hours"], "correct_answer": 2, "explanation": "The current guideline for tPA administration in acute ischemic stroke is within 4.5 hours of symptom onset. This represents the therapeutic window where benefits outweigh risks of hemorrhagic transformation.", "category": "neurology", "year": 3, "difficulty": "2", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        # Endocrinology
        {"id": "endo-001", "question": "A 30-year-old woman presents with headache, palpitations, and sweating episodes. Physical examination reveals a blood pressure of 180/110 mmHg. Which of the following is the most likely diagnosis?", "options": ["Pheochromocytoma", "Cushing syndrome", "Hyperaldosteronism", "Hyperthyroidism", "Essential hypertension"], "correct_answer": 0, "explanation": "The classic triad of headaches, palpitations, and sweating in a hypertensive patient is characteristic of pheochromocytoma. These tumors secrete catecholamines intermittently, leading to paroxysmal symptoms.", "category": "endocrinology", "year": 3, "difficulty": "3", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        {"id": "endo-002", "question": "A 35-year-old woman presents with fatigue, weight gain, and cold intolerance. Laboratory tests show elevated TSH and low free T4. What is the most likely diagnosis?", "options": ["Hyperthyroidism", "Primary hypothyroidism", "Secondary hypothyroidism", "Subclinical hypothyroidism", "Sick euthyroid syndrome"], "correct_answer": 1, "explanation": "Elevated TSH with low free T4 indicates primary hypothyroidism, where the thyroid gland fails to produce adequate thyroid hormone, and the pituitary compensates by increasing TSH secretion.", "category": "endocrinology", "year": 2, "difficulty": "1", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        # Urology
        {"id": "uro-001", "question": "A 60-year-old man presents with painless hematuria. Urine cytology shows atypical cells. Which of the following is the most likely diagnosis?", "options": ["Renal cell carcinoma", "Bladder cancer", "Prostate cancer", "Kidney stones", "Urinary tract infection"], "correct_answer": 1, "explanation": "Painless hematuria in an older adult is a classic presentation of bladder cancer. Urine cytology showing atypical cells further supports this diagnosis.", "category": "urology", "year": 3, "difficulty": "2", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
        {"id": "uro-002", "question": "A 25-year-old man presents with acute onset of severe left flank pain radiating to the groin with nausea. Urinalysis shows microscopic hematuria. What is the most likely diagnosis?", "options": ["Pyelonephritis", "Renal cell carcinoma", "Nephrolithiasis", "Urinary tract infection", "Renal infarction"], "correct_answer": 2, "explanation": "The classic presentation of acute flank pain radiating to the groin with hematuria is characteristic of nephrolithiasis (kidney stones). The pain pattern follows the path of ureteral obstruction.", "category": "urology", "year": 2, "difficulty": "1", "source": "sample", "user_id": None, "verified": True, "tenant_id": tenant_id, "created_at": datetime.utcnow().isoformat()},
    ]
    
    await db.questions.insert_many(comprehensive_questions)
    return {"message": f"Initialized {len(comprehensive_questions)} sample questions"}

# ============================================================================
# UNE PRIORITY QUESTION SYSTEM
# ============================================================================

@api_router.get("/unlock-status")
async def get_unlock_status(user_id: str = Depends(get_current_user)):
    """Get user's question bank unlock status."""
    progress = await db.user_progress.find_one({"user_id": user_id}, {"_id": 0})
    
    if not progress:
        return {
            "full_bank_unlocked": False,
            "qualifying_sessions_completed": 0,
            "sessions_remaining": 3,
            "requirement": "Score 85% or higher on 50+ questions in 3 separate sessions to unlock the full question bank."
        }
    
    full_bank_unlocked = progress.get('full_bank_unlocked', False)
    qualifying_sessions = progress.get('qualifying_sessions_completed', 0)
    
    return {
        "full_bank_unlocked": full_bank_unlocked,
        "qualifying_sessions_completed": qualifying_sessions,
        "sessions_remaining": max(0, 3 - qualifying_sessions),
        "requirement": "Score 85% or higher on 50+ questions in 3 separate sessions to unlock the full question bank."
    }

@api_router.post("/init/une-questions")
async def initialize_une_questions(
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Import UNE priority questions from the Excel file."""
    try:
        une_file_path = ROOT_DIR / 'une_questions.xlsx'
        
        if not une_file_path.exists():
            raise HTTPException(status_code=404, detail="UNE questions file not found")
        
        df = pd.read_excel(une_file_path)
        
        # Clear existing UNE priority questions for this tenant
        deleted = await db.questions.delete_many({"source": "une_priority", "tenant_id": tenant_id})
        logger.info(f"Deleted {deleted.deleted_count} existing UNE priority questions for tenant {tenant_id}")
        
        # Category mapping to handle variations
        category_mapping = {
            'anatomy': 'anatomy',
            'general medicine': 'general medicine',
            'pharmacology': 'pharmacology',
            'microbiology': 'microbiology',
            'pathology': 'pathology',
            'cardiology': 'cardiology',
            'hematology': 'hematology',
            'respiratory': 'respiratory',
            'infectious disease': 'infectious disease',
            'gastroenterology': 'gastroenterology',
            'surgery': 'surgery',
            'endocrinology': 'endocrinology',
            'nephrology': 'nephrology',
            'rheumatology': 'rheumatology',
            'neurology': 'neurology',
            'ophthalmology': 'ophthalmology',
            'psychiatry': 'psychiatry',
            'community medicine': 'community medicine',
            'general': 'general',
        }
        
        questions = []
        skipped = 0
        
        for idx, row in df.iterrows():
            try:
                question_text = str(row['question']).strip()
                if not question_text or question_text == 'nan':
                    skipped += 1
                    continue
                
                # Parse options
                options = []
                for opt_col in ['optionA', 'optionB', 'optionC', 'optionD', 'optionE']:
                    opt_val = row.get(opt_col, '')
                    if pd.notna(opt_val) and str(opt_val).strip() and str(opt_val).strip() != 'nan':
                        options.append(str(opt_val).strip())
                    else:
                        options.append('')  # Keep empty for consistency
                
                # Parse correct answer
                correct_letter = str(row.get('correctAnswer', 'A')).upper().strip()
                if correct_letter in ['A', 'B', 'C', 'D', 'E']:
                    correct_index = ord(correct_letter) - ord('A')
                else:
                    correct_index = 0
                
                # Parse category
                raw_category = str(row.get('category', 'general')).lower().strip()
                category = category_mapping.get(raw_category, 'general')
                
                # Parse other fields
                sub_category = str(row.get('subCategory', '')) if pd.notna(row.get('subCategory')) else None
                explanation = str(row.get('explanation', '')) if pd.notna(row.get('explanation')) else 'No explanation provided.'
                
                year_val = int(row.get('year', 1)) if pd.notna(row.get('year')) else 1
                year_val = max(1, min(6, year_val))
                
                level_val = str(int(row.get('level', 2))) if pd.notna(row.get('level')) else '2'
                
                question_doc = {
                    "id": f"une-{idx+1:05d}",
                    "question": question_text,
                    "options": options,
                    "correct_answer": correct_index,
                    "explanation": explanation,
                    "category": category,
                    "sub_category": sub_category,
                    "year": year_val,
                    "difficulty": level_val,
                    "source": "une_priority",
                    "user_id": None,
                    "verified": True,
                    "quarantined": False,
                    "tenant_id": tenant_id,  # Multi-tenant: assign to current tenant
                    "created_at": datetime.utcnow().isoformat()
                }
                
                questions.append(question_doc)
                
            except Exception as e:
                logger.warning(f"Skipping row {idx} due to error: {e}")
                skipped += 1
                continue
        
        # Insert questions
        if questions:
            await db.questions.insert_many(questions)
        
        return {
            "success": True,
            "questions_imported": len(questions),
            "questions_skipped": skipped,
            "message": f"Successfully imported {len(questions)} UNE priority questions."
        }
        
    except Exception as e:
        logger.error(f"Error importing UNE questions: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/question-bank-stats")
async def get_question_bank_stats(
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Get statistics about question banks for the current tenant."""
    tenant_filter = {
        "$or": [
            {"tenant_id": tenant_id},
            {"tenant_id": {"$exists": False}},
            {"tenant_id": None}
        ]
    }
    
    une_count = await db.questions.count_documents({"source": "une_priority", **tenant_filter})
    total_count = await db.questions.count_documents(tenant_filter)
    other_count = total_count - une_count
    
    return {
        "une_priority_questions": une_count,
        "other_questions": other_count,
        "total_questions": total_count,
        "tenant_id": tenant_id
    }

# ============================================================================
# ADMIN ROUTES (Tenant-Aware)
# ============================================================================

async def verify_admin(user_id: str) -> bool:
    """Verify if user is an admin."""
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    return user and user.get('is_admin', False)

@api_router.get("/admin/users")
async def admin_get_users(
    request: Request,
    tenant_filter_id: Optional[str] = None,
    user_id: str = Depends(get_current_user)
):
    """Get all users with their stats (admin only). Can filter by tenant."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Build query - optionally filter by tenant
    query = {}
    if tenant_filter_id:
        query["tenant_id"] = tenant_filter_id
    
    users = await db.users.find(query, {"_id": 0}).to_list(10000)
    
    # Enrich users with progress data
    enriched_users = []
    for user in users:
        progress = await db.user_progress.find_one({"user_id": user['id']}, {"_id": 0})
        
        # Get session count
        sessions = await db.study_sessions.count_documents({"user_id": user['id']})
        
        enriched_users.append({
            **user,
            "total_questions_answered": progress.get('total_questions_answered', 0) if progress else 0,
            "qualifying_sessions": progress.get('qualifying_sessions_completed', 0) if progress else 0,
            "full_bank_unlocked": progress.get('full_bank_unlocked', False) if progress else False,
            "total_sessions": sessions
        })
    
    return enriched_users

@api_router.get("/admin/stats")
async def admin_get_stats(user_id: str = Depends(get_current_user)):
    """Get platform statistics (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    total_users = await db.users.count_documents({})
    active_subscriptions = await db.users.count_documents({"subscription_status": "active"})
    free_grants = await db.users.count_documents({"subscription_status": "free_grant"})
    
    # Total questions answered across all users
    pipeline = [
        {"$group": {"_id": None, "total": {"$sum": "$total_questions_answered"}}}
    ]
    result = await db.user_progress.aggregate(pipeline).to_list(1)
    total_questions_answered = result[0]['total'] if result else 0
    
    return {
        "total_users": total_users,
        "active_subscriptions": active_subscriptions,
        "free_grants": free_grants,
        "total_questions_answered": total_questions_answered
    }

@api_router.post("/admin/grant-access")
async def admin_grant_access(
    data: dict,
    user_id: str = Depends(get_current_user)
):
    """Grant free subscription access to a user (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    target_user_id = data.get('user_id')
    subscription_plan = data.get('subscription_plan', 'monthly')
    duration_days = data.get('duration_days', 30)
    
    if not target_user_id:
        raise HTTPException(status_code=400, detail="user_id is required")
    
    # Calculate end date
    end_date = datetime.utcnow() + timedelta(days=duration_days)
    
    # AI limits based on subscription (50% of cost budget)
    ai_limits = {
        'weekly': 0,
        'monthly': 5,
        'quarterly': 10,
        'annual': 10
    }
    
    # Get user's current subscription for history
    target_user = await db.users.find_one({"id": target_user_id}, {"_id": 0})
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Log subscription history
    await db.subscription_history.insert_one({
        "user_id": target_user_id,
        "action": "upgrade",
        "from_plan": target_user.get('subscription_plan') or 'free',
        "to_plan": subscription_plan,
        "granted_by": user_id,
        "grant_type": "free_grant",
        "start_date": datetime.utcnow().isoformat(),
        "end_date": end_date.isoformat(),
        "duration_days": duration_days,
        "timestamp": datetime.utcnow().isoformat()
    })
    
    # Update user
    await db.users.update_one(
        {"id": target_user_id},
        {"$set": {
            "subscription_status": "free_grant",
            "subscription_plan": subscription_plan,
            "subscription_tier": subscription_plan,  # Also set tier for consistency
            "subscription_start": datetime.utcnow().isoformat(),
            "subscription_end": end_date.isoformat(),
            "ai_max_daily_uses": ai_limits.get(subscription_plan, 10)
        }}
    )
    
    return {"success": True, "message": f"Granted {subscription_plan} access until {end_date.date()}"}

@api_router.post("/admin/revoke-access")
async def admin_revoke_access(
    data: dict,
    user_id: str = Depends(get_current_user)
):
    """Revoke subscription access from a user (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    target_user_id = data.get('user_id')
    
    if not target_user_id:
        raise HTTPException(status_code=400, detail="user_id is required")
    
    # Get user's current subscription for history
    target_user = await db.users.find_one({"id": target_user_id}, {"_id": 0})
    if target_user and target_user.get('subscription_plan'):
        # Log subscription history
        await db.subscription_history.insert_one({
            "user_id": target_user_id,
            "action": "revoke",
            "from_plan": target_user.get('subscription_plan'),
            "to_plan": "free",
            "revoked_by": user_id,
            "timestamp": datetime.utcnow().isoformat()
        })
    
    # Update user
    await db.users.update_one(
        {"id": target_user_id},
        {"$set": {
            "subscription_status": "free",
            "subscription_plan": None,
            "subscription_tier": "free",
            "subscription_end": None,
            "ai_max_daily_uses": 0
        }}
    )
    
    return {"success": True, "message": "Access revoked"}

@api_router.post("/admin/toggle-admin")
async def admin_toggle_admin(
    data: dict,
    user_id: str = Depends(get_current_user)
):
    """Toggle admin status for a user (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    target_user_id = data.get('user_id')
    is_admin = data.get('is_admin', False)
    
    if not target_user_id:
        raise HTTPException(status_code=400, detail="user_id is required")
    
    # Prevent removing own admin access
    if target_user_id == user_id and not is_admin:
        raise HTTPException(status_code=400, detail="Cannot remove your own admin access")
    
    # Update user
    await db.users.update_one(
        {"id": target_user_id},
        {"$set": {"is_admin": is_admin}}
    )
    
    return {"success": True, "message": f"Admin access {'granted' if is_admin else 'revoked'}"}

@api_router.post("/admin/bootstrap")
async def bootstrap_admin(
    data: dict,
    request: Request,
    tenant_id: str = Depends(get_current_tenant)
):
    """Bootstrap first admin user using secret key.
    
    This endpoint allows promoting a user to admin using a secret bootstrap key.
    Used for initial setup when no admin exists yet.
    """
    email = data.get('email')
    bootstrap_key = data.get('bootstrap_key')
    
    # Secret bootstrap key - should match environment variable or default
    expected_key = os.environ.get('ADMIN_BOOTSTRAP_KEY', 'medmcq-admin-bootstrap-2024')
    
    if not email or not bootstrap_key:
        raise HTTPException(status_code=400, detail="email and bootstrap_key are required")
    
    if bootstrap_key != expected_key:
        raise HTTPException(status_code=403, detail="Invalid bootstrap key")
    
    # Find user with tenant context
    user = await db.users.find_one({"email": email, "tenant_id": tenant_id})
    
    if not user:
        # Try without tenant filter as fallback
        user = await db.users.find_one({"email": email})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
    
    # Force update regardless of current state
    result = await db.users.update_one(
        {"id": user["id"]},
        {"$set": {"is_admin": True, "email_verified": True}}
    )
    
    # Verify the update worked
    updated_user = await db.users.find_one({"id": user["id"]})
    
    logger.info(f"Bootstrap: User {email} (id: {user['id']}, tenant: {user.get('tenant_id')}) - is_admin now: {updated_user.get('is_admin')}")
    
    return {
        "success": True, 
        "message": f"User {email} has been granted admin access",
        "user_id": user["id"],
        "is_admin": updated_user.get('is_admin'),
        "modified_count": result.modified_count
    }

@api_router.get("/admin/user/{target_user_id}/subscription-history")
async def admin_get_subscription_history(
    target_user_id: str,
    user_id: str = Depends(get_current_user)
):
    """Get subscription history for a specific user (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Get user details
    user = await db.users.find_one({"id": target_user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Get subscription history
    history = await db.subscription_history.find(
        {"user_id": target_user_id}, 
        {"_id": 0}
    ).sort("timestamp", -1).to_list(100)
    
    return {
        "user": {
            "id": user.get("id"),
            "email": user.get("email"),
            "full_name": user.get("full_name"),
            "current_plan": user.get("subscription_plan") or "free",
            "subscription_status": user.get("subscription_status") or "free",
            "subscription_start": user.get("subscription_start"),
            "subscription_end": user.get("subscription_end"),
            "email_verified": user.get("email_verified", False),
            "marketing_consent": user.get("marketing_consent", False),
            "created_at": user.get("created_at")
        },
        "history": history
    }

# ============================================================================
# CRM ROUTES (Admin Only)
# ============================================================================

@api_router.get("/crm/contacts")
async def crm_get_contacts(user_id: str = Depends(get_current_user)):
    """Get all CRM contacts (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    contacts = await db.crm_contacts.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return contacts

@api_router.get("/crm/contacts/{contact_id}")
async def crm_get_contact(contact_id: str, user_id: str = Depends(get_current_user)):
    """Get single CRM contact with details (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    contact = await db.crm_contacts.find_one({"id": contact_id}, {"_id": 0})
    if not contact:
        raise HTTPException(status_code=404, detail="Contact not found")
    return contact

@api_router.post("/crm/contacts")
async def crm_create_contact(data: dict, user_id: str = Depends(get_current_user)):
    """Create a new CRM contact (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    contact = {
        "id": str(uuid.uuid4()),
        "name": data.get("name"),
        "email": data.get("email"),
        "phone": data.get("phone"),
        "company": data.get("company"),
        "status": data.get("status", "lead"),
        "source": data.get("source"),
        "notes_history": [],
        "created_at": datetime.utcnow().isoformat(),
        "updated_at": datetime.utcnow().isoformat(),
        "created_by": user_id
    }
    
    # Add initial note if provided
    if data.get("notes"):
        contact["notes_history"].append({
            "content": data["notes"],
            "created_at": datetime.utcnow().isoformat(),
            "created_by": user_id
        })
    
    await db.crm_contacts.insert_one(contact)
    return contact

@api_router.put("/crm/contacts/{contact_id}")
async def crm_update_contact(contact_id: str, data: dict, user_id: str = Depends(get_current_user)):
    """Update a CRM contact (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    update_data = {k: v for k, v in data.items() if k not in ["id", "created_at", "notes_history"]}
    update_data["updated_at"] = datetime.utcnow().isoformat()
    
    await db.crm_contacts.update_one(
        {"id": contact_id},
        {"$set": update_data}
    )
    
    return {"success": True}

@api_router.delete("/crm/contacts/{contact_id}")
async def crm_delete_contact(contact_id: str, user_id: str = Depends(get_current_user)):
    """Delete a CRM contact (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    await db.crm_contacts.delete_one({"id": contact_id})
    # Also delete related tasks
    await db.crm_tasks.delete_many({"contact_id": contact_id})
    
    return {"success": True}

@api_router.post("/crm/contacts/{contact_id}/notes")
async def crm_add_note(contact_id: str, data: dict, user_id: str = Depends(get_current_user)):
    """Add a note to a CRM contact (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    note = {
        "content": data.get("content"),
        "created_at": datetime.utcnow().isoformat(),
        "created_by": user_id
    }
    
    await db.crm_contacts.update_one(
        {"id": contact_id},
        {
            "$push": {"notes_history": note},
            "$set": {"updated_at": datetime.utcnow().isoformat()}
        }
    )
    
    return {"success": True}

@api_router.get("/crm/tasks")
async def crm_get_tasks(user_id: str = Depends(get_current_user)):
    """Get all CRM tasks (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    tasks = await db.crm_tasks.find({}, {"_id": 0}).sort("due_date", 1).to_list(1000)
    
    # Enrich with contact names
    for task in tasks:
        if task.get("contact_id"):
            contact = await db.crm_contacts.find_one({"id": task["contact_id"]}, {"name": 1})
            task["contact_name"] = contact.get("name") if contact else None
    
    return tasks

@api_router.post("/crm/tasks")
async def crm_create_task(data: dict, user_id: str = Depends(get_current_user)):
    """Create a new CRM task (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    task = {
        "id": str(uuid.uuid4()),
        "title": data.get("title"),
        "description": data.get("description"),
        "due_date": data.get("due_date"),
        "priority": data.get("priority", "medium"),
        "contact_id": data.get("contact_id"),
        "completed": False,
        "created_at": datetime.utcnow().isoformat(),
        "created_by": user_id
    }
    
    await db.crm_tasks.insert_one(task)
    return task

@api_router.put("/crm/tasks/{task_id}")
async def crm_update_task(task_id: str, data: dict, user_id: str = Depends(get_current_user)):
    """Update a CRM task (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    update_data = {k: v for k, v in data.items() if k not in ["id", "created_at"]}
    if data.get("completed"):
        update_data["completed_at"] = datetime.utcnow().isoformat()
    
    await db.crm_tasks.update_one(
        {"id": task_id},
        {"$set": update_data}
    )
    
    return {"success": True}

@api_router.get("/crm/stats")
async def crm_get_stats(user_id: str = Depends(get_current_user)):
    """Get CRM statistics (admin only)."""
    if not await verify_admin(user_id):
        raise HTTPException(status_code=403, detail="Admin access required")
    
    total_contacts = await db.crm_contacts.count_documents({})
    leads = await db.crm_contacts.count_documents({"status": "lead"})
    prospects = await db.crm_contacts.count_documents({"status": "prospect"})
    customers = await db.crm_contacts.count_documents({"status": "customer"})
    pending_tasks = await db.crm_tasks.count_documents({"completed": False})
    
    return {
        "total_contacts": total_contacts,
        "leads": leads,
        "prospects": prospects,
        "customers": customers,
        "pending_tasks": pending_tasks
    }

# ============================================================================
# BASIC ROUTES
# ============================================================================

@api_router.get("/")
async def root():
    return {"message": "MedMCQ API - Medical Student Learning Platform"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}

# ============================================================================
# STRIPE PAYMENT ROUTES (DISABLED)
# ============================================================================

# Payment gateway is temporarily disabled
# Uncomment when ready to enable payments

# from emergentintegrations.payments.stripe.checkout import StripeCheckout, CheckoutSessionResponse, CheckoutStatusResponse, CheckoutSessionRequest
from fastapi import Request

# Define subscription packages (amounts in dollars) - kept for reference
SUBSCRIPTION_PACKAGES = {
    "weekly": {"name": "Weekly", "amount": 9.99, "period": "week"},
    "monthly": {"name": "Monthly", "amount": 29.99, "period": "month"},
    "quarterly": {"name": "Quarterly", "amount": 79.99, "period": "3 months"},
    "annual": {"name": "Annual", "amount": 249.99, "period": "year"}
}

@api_router.post("/payments/create-checkout")
async def create_checkout_session(
    request: Request,
    package_id: str,
    user_id: str = Depends(get_current_user)
):
    """Create a Stripe checkout session for subscription - DISABLED"""
    raise HTTPException(
        status_code=503,
        detail="Payment system is temporarily disabled. Please check back later."
    )

@api_router.get("/payments/status/{session_id}")
async def get_payment_status(
    session_id: str,
    request: Request,
    user_id: str = Depends(get_current_user)
):
    """Get payment status for a checkout session - DISABLED"""
    raise HTTPException(
        status_code=503,
        detail="Payment system is temporarily disabled. Please check back later."
    )

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhooks - DISABLED"""
    return {"status": "disabled", "message": "Payment system is temporarily disabled."}

@api_router.get("/payments/config")
async def get_stripe_config():
    """Get Stripe publishable key for frontend - DISABLED"""
    return {"publishable_key": None, "disabled": True, "message": "Payment system is temporarily disabled."}

@api_router.get("/user/subscription")
async def get_user_subscription(user_id: str = Depends(get_current_user)):
    """Get user's subscription status"""
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {
        "subscription_status": user.get("subscription_status", "free"),
        "subscription_plan": user.get("subscription_plan"),
        "subscription_end": user.get("subscription_end"),
        "payments_disabled": True
    }


# ============================================================================
# CONTACT FORM ROUTE
# ============================================================================

@api_router.post("/contact")
async def submit_contact_form(data: dict):
    """Submit contact form - sends email to support."""
    name = data.get("name")
    email = data.get("email")
    subject = data.get("subject", "Contact Form Submission")
    message = data.get("message")
    
    if not all([name, email, message]):
        raise HTTPException(status_code=400, detail="Name, email, and message are required")
    
    # Store in database
    contact_entry = {
        "id": str(uuid.uuid4()),
        "name": name,
        "email": email,
        "subject": subject,
        "message": message,
        "submitted_at": datetime.utcnow().isoformat(),
        "status": "new"
    }
    await db.contact_submissions.insert_one(contact_entry)
    
    # Send email notification to support
    email_sent = send_contact_form_notification(
        sender_name=name,
        sender_email=email,
        subject=subject,
        message=message
    )
    
    if email_sent:
        logger.info(f"Contact form submitted by {email}, notification sent to support")
    else:
        logger.warning(f"Contact form submitted by {email}, but notification email failed")
    
    return {"success": True, "message": "Thank you for your message! We will get back to you within 24 hours."}


# Include router in app
app.include_router(api_router)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

@api_router.post("/init/comprehensive-sample-data")
async def initialize_comprehensive_sample_data():
    """Initialize comprehensive sample questions for better demo"""
    # Clear existing sample data
    await db.questions.delete_many({"source": "sample"})
    
    comprehensive_questions = [
        # Respiratory Questions
        {
            "id": "resp-001",
            "question": "A 65-year-old man with a 40-pack-year smoking history presents with worsening dyspnea and a chronic cough. Pulmonary function tests show a decreased FEV1/FVC ratio, increased total lung capacity, and decreased diffusing capacity. Which of the following is the most likely diagnosis?",
            "options": ["Asthma", "Chronic bronchitis", "Emphysema", "Bronchiectasis"],
            "correct_answer": 2,
            "explanation": "The patient's presentation and pulmonary function tests are classic for emphysema. Key findings include: 1) Decreased FEV1/FVC ratio indicating obstructive lung disease, 2) Increased total lung capacity indicating hyperinflation, and 3) Decreased diffusing capacity due to destruction of alveolar walls.",
            "category": "respiratory",
            "year": 2,
            "difficulty": "2",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        {
            "id": "resp-002",
            "question": "A 28-year-old woman presents to the emergency department with sudden onset of dyspnea and right-sided pleuritic chest pain. She recently returned from a long international flight. Her vital signs show tachycardia and mild hypoxemia. What is the most appropriate initial diagnostic test?",
            "options": ["Chest X-ray", "CT pulmonary angiography", "D-dimer", "Ventilation-perfusion scan"],
            "correct_answer": 2,
            "explanation": "Given the clinical suspicion for pulmonary embolism (recent travel, pleuritic chest pain, dyspnea), D-dimer is the most appropriate initial test. A negative D-dimer can rule out PE in low-risk patients. If positive, proceed to CT pulmonary angiography for definitive diagnosis.",
            "category": "respiratory",
            "year": 3,
            "difficulty": "2",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        # Cardiology Questions
        {
            "id": "cardio-001",
            "question": "A 55-year-old woman presents with crushing substernal chest pain that radiates to her left arm. ECG shows ST-segment elevation in leads II, III, and aVF. Which coronary artery is most likely occluded?",
            "options": ["Left anterior descending artery", "Left circumflex artery", "Right coronary artery", "Left main coronary artery"],
            "correct_answer": 2,
            "explanation": "ST-segment elevation in leads II, III, and aVF indicates an inferior wall myocardial infarction. The right coronary artery (RCA) supplies the inferior wall of the left ventricle in 85% of people.",
            "category": "cardiology",
            "year": 3,
            "difficulty": "3",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        {
            "id": "cardio-002",
            "question": "A 45-year-old man with hypertension presents for a routine check-up. His blood pressure is 145/95 mmHg. Which of the following is the most appropriate first-line antihypertensive medication for this patient without other comorbidities?",
            "options": ["Thiazide diuretic", "Beta-blocker", "Calcium channel blocker", "Alpha-blocker"],
            "correct_answer": 0,
            "explanation": "Thiazide diuretics are recommended as first-line therapy for uncomplicated hypertension according to JNC guidelines. They have been shown to reduce cardiovascular events and are cost-effective.",
            "category": "cardiology",
            "year": 2,
            "difficulty": "1",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        # Neurology Questions
        {
            "id": "neuro-001",
            "question": "A 45-year-old man with chronic alcohol use presents with confusion, ataxia, and nystagmus. Which vitamin deficiency is most likely responsible?",
            "options": ["Vitamin B1 (Thiamine)", "Vitamin B12", "Vitamin C", "Vitamin D"],
            "correct_answer": 0,
            "explanation": "Wernicke's encephalopathy, characterized by the triad of confusion, ataxia, and nystagmus, is caused by thiamine (Vitamin B1) deficiency, commonly seen in chronic alcoholics.",
            "category": "neurology",
            "year": 2,
            "difficulty": "2",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        {
            "id": "neuro-002",
            "question": "A 72-year-old woman presents with sudden onset of right-sided weakness and inability to speak. CT scan of the head shows no hemorrhage. What is the time window for administering tissue plasminogen activator (tPA)?",
            "options": ["1.5 hours", "3 hours", "4.5 hours", "6 hours"],
            "correct_answer": 2,
            "explanation": "The current guideline for tPA administration in acute ischemic stroke is within 4.5 hours of symptom onset. This represents the therapeutic window where benefits outweigh risks of hemorrhagic transformation.",
            "category": "neurology",
            "year": 3,
            "difficulty": "2",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        # Endocrinology Questions
        {
            "id": "endo-001",
            "question": "A 30-year-old woman presents with headache, palpitations, and sweating episodes. Physical examination reveals a blood pressure of 180/110 mmHg. Which of the following is the most likely diagnosis?",
            "options": ["Pheochromocytoma", "Cushing syndrome", "Hyperaldosteronism", "Hyperthyroidism"],
            "correct_answer": 0,
            "explanation": "The classic triad of headaches, palpitations, and sweating in a hypertensive patient is characteristic of pheochromocytoma. These tumors secrete catecholamines intermittently, leading to paroxysmal symptoms.",
            "category": "endocrinology",
            "year": 3,
            "difficulty": "3",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        {
            "id": "endo-002",
            "question": "A 35-year-old woman presents with fatigue, weight gain, and cold intolerance. Laboratory tests show elevated TSH and low free T4. What is the most likely diagnosis?",
            "options": ["Hyperthyroidism", "Primary hypothyroidism", "Secondary hypothyroidism", "Subclinical hypothyroidism"],
            "correct_answer": 1,
            "explanation": "Elevated TSH with low free T4 indicates primary hypothyroidism, where the thyroid gland fails to produce adequate thyroid hormone, and the pituitary compensates by increasing TSH secretion.",
            "category": "endocrinology",
            "year": 2,
            "difficulty": "1",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        # Urology Questions
        {
            "id": "uro-001",
            "question": "A 60-year-old man presents with painless hematuria. Urine cytology shows atypical cells. Which of the following is the most likely diagnosis?",
            "options": ["Renal cell carcinoma", "Bladder cancer", "Prostate cancer", "Kidney stones"],
            "correct_answer": 1,
            "explanation": "Painless hematuria in an older adult is a classic presentation of bladder cancer. Urine cytology showing atypical cells further supports this diagnosis.",
            "category": "urology",
            "year": 3,
            "difficulty": "2",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
        {
            "id": "uro-002",
            "question": "A 25-year-old man presents with acute onset of severe left flank pain radiating to the groin with nausea. Urinalysis shows microscopic hematuria. What is the most likely diagnosis?",
            "options": ["Pyelonephritis", "Renal cell carcinoma", "Nephrolithiasis", "Urinary tract infection"],
            "correct_answer": 2,
            "explanation": "The classic presentation of acute flank pain radiating to the groin with hematuria is characteristic of nephrolithiasis (kidney stones). The pain pattern follows the path of ureteral obstruction.",
            "category": "urology",
            "year": 2,
            "difficulty": "1",
            "source": "sample",
            "user_id": None,
            "verified": True,
            "created_at": datetime.utcnow().isoformat()
        },
    ]
    
    await db.questions.insert_many(comprehensive_questions)
    
    return {"message": f"Initialized {len(comprehensive_questions)} comprehensive sample questions"}
