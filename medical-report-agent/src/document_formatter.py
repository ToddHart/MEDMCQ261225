"""
Document Formatter - Creates professionally formatted reports with graphs
Generates Word and PDF documents with proper formatting, fonts, and visualizations
"""

import os
from pathlib import Path
from datetime import datetime
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠ python-docx not installed - Word document generation unavailable")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠ reportlab not installed - PDF generation unavailable")

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Non-interactive backend
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("⚠ matplotlib not installed - graph generation unavailable")

import json


class DocumentFormatter:
    """Formats reports into professional documents with graphs"""

    def __init__(self, output_dir="output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    def create_test_graphs(self, patient_data, output_path):
        """Generate graphs showing test results"""
        if not MATPLOTLIB_AVAILABLE:
            print("⚠ Matplotlib not available - skipping graphs")
            return None

        # Collect scores from all tests
        test_names = []
        percentiles = []
        scores_data = []

        for test in patient_data['tests']:
            if 'percentile' in test:
                test_names.append(test['test_code'])
                percentiles.append(test['percentile'])
            elif 'scores' in test and isinstance(test['scores'], dict):
                # Multi-score test
                for score_name, score_value in test['scores'].items():
                    test_names.append(f"{test['test_code']}\n{score_name[:10]}")
                    percentiles.append(score_value)

        if not percentiles:
            print("⚠ No percentile data found for graphing")
            return None

        # Create figure with subplots
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

        # Bar chart of percentiles
        colors = ['#2E86AB' if p >= 50 else '#A23B72' for p in percentiles]
        ax1.bar(range(len(test_names)), percentiles, color=colors, alpha=0.7)
        ax1.axhline(y=50, color='gray', linestyle='--', linewidth=1, label='Average (50th %ile)')
        ax1.set_xlabel('Tests', fontsize=11, fontweight='bold')
        ax1.set_ylabel('Percentile Score', fontsize=11, fontweight='bold')
        ax1.set_title('Test Performance Overview', fontsize=13, fontweight='bold')
        ax1.set_xticks(range(len(test_names)))
        ax1.set_xticklabels(test_names, rotation=45, ha='right', fontsize=8)
        ax1.set_ylim(0, 100)
        ax1.legend()
        ax1.grid(axis='y', alpha=0.3)

        # Performance category pie chart
        categories = {'Above Average (>75)': 0, 'Average (25-75)': 0, 'Below Average (<25)': 0}
        for p in percentiles:
            if p > 75:
                categories['Above Average (>75)'] += 1
            elif p >= 25:
                categories['Average (25-75)'] += 1
            else:
                categories['Below Average (<25)'] += 1

        # Filter out zero values
        filtered_categories = {k: v for k, v in categories.items() if v > 0}

        if filtered_categories:
            ax2.pie(filtered_categories.values(), labels=filtered_categories.keys(),
                   autopct='%1.0f%%', startangle=90,
                   colors=['#06A77D', '#2E86AB', '#A23B72'])
            ax2.set_title('Performance Distribution', fontsize=13, fontweight='bold')

        plt.tight_layout()
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()

        print(f"✓ Generated test results graph: {output_path}")
        return output_path

    def create_word_document(self, report_data, patient_data):
        """Create a formatted Word document"""
        if not DOCX_AVAILABLE:
            print("⚠ Word document generation not available")
            return None

        print("\nCreating Word document...")

        doc = Document()

        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Title
        title = doc.add_heading('PSYCHOLOGICAL ASSESSMENT REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        # Add metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Patient ID: {report_data['patient_id']}").bold = True
        meta.add_run(f"\nGenerated: {report_data['generated_date']}")

        doc.add_paragraph()

        # Add graph if available
        graph_path = self.output_dir / f"graph_{report_data['patient_id']}.png"
        if self.create_test_graphs(patient_data, graph_path):
            doc.add_picture(str(graph_path), width=Inches(6))
            doc.add_paragraph()

        # Add report content
        # Split content into paragraphs and format
        content_lines = report_data['content'].split('\n')

        for line in content_lines:
            if not line.strip():
                continue

            # Check if line is a header (all caps or ends with colon)
            if line.isupper() and len(line.strip()) > 3:
                # Main header
                heading = doc.add_heading(line.strip(), level=1)
                heading_run = heading.runs[0]
                heading_run.font.size = Pt(13)
                heading_run.font.color.rgb = RGBColor(0, 51, 102)
            elif line.strip().endswith(':') and line.strip()[0].isupper():
                # Subheading
                subheading = doc.add_paragraph(line.strip())
                subheading_run = subheading.runs[0]
                subheading_run.bold = True
                subheading_run.font.size = Pt(11)
            elif line.strip().startswith(('•', '-', '▸', '★', '✓', '→')):
                # Bullet point
                doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
            elif line.strip()[0].isdigit() and line.strip()[1] in '.):':
                # Numbered list
                doc.add_paragraph(line.strip()[2:].strip(), style='List Number')
            else:
                # Regular paragraph
                doc.add_paragraph(line.strip())

        # Save document
        output_path = self.output_dir / f"report_{report_data['patient_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(output_path)

        print(f"✓ Word document saved: {output_path}")
        return output_path

    def create_pdf_document(self, report_data, patient_data):
        """Create a formatted PDF document"""
        if not PDF_AVAILABLE:
            print("⚠ PDF generation not available")
            return None

        print("\nCreating PDF document...")

        output_path = self.output_dir / f"report_{report_data['patient_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

        # Create PDF
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='#003366',
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=13,
            textColor='#003366',
            spaceAfter=10,
            spaceBefore=10,
            fontName='Helvetica-Bold'
        )

        # Title
        story.append(Paragraph("PSYCHOLOGICAL ASSESSMENT REPORT", title_style))
        story.append(Spacer(1, 0.2*inch))

        # Metadata
        story.append(Paragraph(f"<b>Patient ID:</b> {report_data['patient_id']}", styles['Normal']))
        story.append(Paragraph(f"<b>Generated:</b> {report_data['generated_date']}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))

        # Add graph
        graph_path = self.output_dir / f"graph_{report_data['patient_id']}.png"
        if self.create_test_graphs(patient_data, graph_path):
            img = Image(str(graph_path), width=6*inch, height=2.5*inch)
            story.append(img)
            story.append(Spacer(1, 0.2*inch))

        # Add content
        content_lines = report_data['content'].split('\n')

        for line in content_lines:
            if not line.strip():
                story.append(Spacer(1, 0.1*inch))
                continue

            if line.isupper() and len(line.strip()) > 3:
                story.append(Paragraph(line.strip(), heading_style))
            elif line.strip().endswith(':') and line.strip()[0].isupper():
                story.append(Paragraph(f"<b>{line.strip()}</b>", styles['Normal']))
            else:
                story.append(Paragraph(line.strip(), styles['Normal']))

        doc.build(story)

        print(f"✓ PDF document saved: {output_path}")
        return output_path

    def format_report(self, report_data, patient_data, format='both'):
        """Format report in specified format(s)"""
        print(f"\n{'='*60}")
        print("DOCUMENT FORMATTING")
        print(f"{'='*60}")

        outputs = {}

        if format in ['word', 'both']:
            word_path = self.create_word_document(report_data, patient_data)
            if word_path:
                outputs['word'] = word_path

        if format in ['pdf', 'both']:
            pdf_path = self.create_pdf_document(report_data, patient_data)
            if pdf_path:
                outputs['pdf'] = pdf_path

        print(f"\n{'='*60}")
        print("Document Formatting Complete!")
        print(f"{'='*60}")

        return outputs


if __name__ == "__main__":
    # Test document formatting
    # Load sample patient data
    with open("../data/patient_db/patients.json", 'r') as f:
        patients = json.load(f)
        patient_data = patients[0]

    # Sample report data
    report_data = {
        'patient_id': patient_data['patient_id'],
        'patient_name': patient_data['name'],
        'generated_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'content': "PSYCHOLOGICAL ASSESSMENT REPORT\n\nReason for Referral:\nTest content here."
    }

    formatter = DocumentFormatter("../output")
    formatter.format_report(report_data, patient_data)
